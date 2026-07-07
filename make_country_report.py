"""
make_country_report.py
======================
指定した国の概況資料（A4 MS Word）を自動生成するツール。

【構成】
  上部右  : 外務省 国・地域HP から国旗・周辺地図を取得して貼付
  上部左  : 基礎データ（国名・面積・人口・首都・民族・言語・宗教）
  中段    : 経済状況（GDP・貿易等）
  下段    : 日本との関係（要人往来・条約・貿易・日系企業・在留邦人・大使）

出力先: C:\\Users\\shondo\\Desktop\\agent_project\\country_reports
依存:   pip install anthropic python-docx pillow beautifulsoup4
"""

import io, json, re, sys, time, urllib.request, urllib.parse
from io import BytesIO
from pathlib import Path
from datetime import datetime

if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
    sys.stdin  = io.TextIOWrapper(sys.stdin.buffer,  encoding="utf-8", errors="replace")

try:
    import anthropic
except ImportError:
    sys.exit("anthropic が未インストールです。pip install anthropic を実行してください。")

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("python-docx が未インストールです。pip install python-docx を実行してください。")

try:
    from PIL import Image as PILImage
except ImportError:
    PILImage = None

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

# ── 設定 ─────────────────────────────────────────────────────────────────────
CLAUDE_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
CLAUDE_MODEL   = "claude-sonnet-4-6"
OUT_DIR        = Path(r"C:\Users\shondo\Desktop\agent_project\country_reports")
IMG_DIR        = OUT_DIR / "images"
F              = 11   # 本文フォントサイズ (pt)
HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept": (
        "text/html,application/xhtml+xml,application/xml;q=0.9,"
        "image/avif,image/webp,*/*;q=0.8"
    ),
    "Accept-Language": "ja,en-US;q=0.9,en;q=0.8",
    "Accept-Encoding": "gzip, deflate, br",
    "Referer": "https://www.mofa.go.jp/",
    "Connection": "keep-alive",
}
IMG_HEADERS = {          # 画像ダウンロード用（Accept を画像向けに変更）
    "User-Agent": HEADERS["User-Agent"],
    "Accept": "image/avif,image/webp,image/apng,image/*,*/*;q=0.8",
    "Accept-Language": "ja,en-US;q=0.9",
    "Referer": "https://www.mofa.go.jp/",
    "Connection": "keep-alive",
}
MOFA_BASE = "https://www.mofa.go.jp"

# ── Claude プロンプト ──────────────────────────────────────────────────────────

# ① 外務省 URL ＋ 基礎データ
BASIC_DATA_PROMPT = """\
あなたは日本の外務省（MOFA）の国・地域情報に精通した専門家です。
以下の国について、外務省「国・地域」HPの情報をもとに基礎データを提供してください。

【対象国】{country}

【出力ルール】
- JSON のみを出力（前置き・コードブロック不要）
- 情報が不明な場合は "不明" と記載
- mofa_page_url は https://www.mofa.go.jp/mofaj/area/ 配下の正確なURL

{{
  "country_name_ja": "正式国名（日本語）",
  "country_name_en": "正式国名（英語）",
  "area_km2": "面積（数値＋単位、例: 約983万km²）",
  "population": "人口（例: 約3億3,000万人（2023年））",
  "capital": "首都名",
  "ethnic_groups": ["主要民族（各25文字以内、3件程度）"],
  "languages": ["公用語・主要言語（各20文字以内）"],
  "religions": ["主要宗教（各20文字以内）"],
  "mofa_page_url": "外務省の当該国ページURL",
  "mofa_flag_img_url": "外務省ページ内の国旗画像URL（推定できれば記載、不明なら空文字）",
  "mofa_map_img_url": "外務省ページ内の周辺地図画像URL（推定できれば記載、不明なら空文字）",
  "wikipedia_flag_url": "Wikipedia等で公開されている国旗SVG/PNG画像の直接URL"
}}
"""

# ② 経済状況
ECONOMIC_DATA_PROMPT = """\
あなたは国際経済の専門家です。
以下の国の経済状況について、最新の公的データ（IMF・世界銀行・外務省等）をもとに情報を提供してください。

【対象国】{country}（{country_en}）

【出力ルール】
- JSON のみを出力（前置き・コードブロック不要）
- 不明・該当なしの場合は "不明" と記載
- データは最新年度を優先し、年度を明記すること（例: 2023年）

{{
  "major_industries": ["主要産業（各30文字以内、5件程度）"],
  "gdp": "名目GDP（例: 約27.4兆ドル（2023年））",
  "gdp_per_capita": "一人当たりGDP（例: 約8.2万ドル（2023年））",
  "gdp_growth_rate": "経済成長率（例: 2.5%（2023年））",
  "inflation_rate": "物価上昇率（例: 3.4%（2023年））",
  "unemployment_rate": "失業率（例: 3.7%（2023年））",
  "total_exports": "総輸出額（例: 約3.1兆ドル（2022年））",
  "total_imports": "総輸入額（例: 約3.4兆ドル（2022年））",
  "major_export_items": ["主要輸出品目（各25文字以内、5件程度）"],
  "major_import_items": ["主要輸入品目（各25文字以内、5件程度）"],
  "major_trade_partners": ["主要貿易相手国（各20文字以内、5件程度）"],
  "data_source": "データ出典（例: IMF World Economic Outlook 2024年）"
}}
"""

# ③ 日本との関係
JAPAN_RELATIONS_PROMPT = """\
あなたは日本の外交・通商関係に精通した専門家です。
以下の国と日本との関係について、外務省情報・貿易統計等をもとに情報を提供してください。

【対象国】{country}（{country_en}）

【出力ルール】
- JSON のみを出力（前置き・コードブロック不要）
- 不明・該当なしの場合は "不明" または空リスト []
- 要人往来は直近2〜3年（{recent_years}年〜現在）の情報を優先

{{
  "vip_visits": [
    {{
      "date": "20XX年XX月",
      "summary": "往来内容（例: ○○首相が日本訪問・首脳会談）（40文字以内）"
    }}
  ],
  "treaties": [
    {{
      "name": "条約・協定名（30文字以内）",
      "year": "締結年（例: 2023年）"
    }}
  ],
  "japan_exports_to": "日本からの輸出額（例: 約1.5兆円（2022年度））",
  "japan_imports_from": "日本への輸入額（例: 約2.1兆円（2022年度））",
  "major_japan_export_items": ["日本の主要輸出品目（各25文字以内、5件）"],
  "major_japan_import_items": ["日本の主要輸入品目（各25文字以内、5件）"],
  "japanese_companies_count": "日系企業進出数（例: 約1,800社（2023年））",
  "japanese_companies_notable": ["主な日系進出企業（各20文字以内、5〜8社）"],
  "japanese_residents": "在留邦人数（例: 約4.5万人（2023年10月）、外務省統計）",
  "japan_ambassador": "駐{country}日本大使名（例: 山田太郎（2024年着任））",
  "country_ambassador_to_japan": "駐日{country}大使名（例: ○○（2023年着任））"
}}
"""


# ── XML / フォントヘルパー ─────────────────────────────────────────────────────
def _set_ea_font(run, font_name="MS明朝"):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _remove_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:color"), "FFFFFF")   # 白色枠線
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _set_cell_w(cell, cm):
    tc   = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr"); tc.insert(0, tcPr)
    tcW  = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
    tcW.set(qn("w:w"),    str(int(cm * 567)))
    tcW.set(qn("w:type"), "dxa")


def _add_run(para_obj, text, size=F, bold=False, color=None, italic=False):
    run = para_obj.add_run(text)
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.name   = "MS明朝"
    _set_ea_font(run)
    if color:
        run.font.color.rgb = color
    return run


def _para(doc, text="", size=F, bold=False, color=None,
          align=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=0, space_after=3, underline=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        r = _add_run(p, text, size=size, bold=bold, color=color)
        r.font.underline = underline
    return p


def _section_title(doc, text):
    """セクションヘッダー（太字＋下線、make_cv_plus.py と同スタイル）"""
    return _para(doc, text, size=14, bold=True, underline=True,
                 space_before=40, space_after=10)


def _kv_row(doc, label, value, label_w=4.5, value_w=12.0):
    """1行 Key-Value テーブル行を直接追加（テーブルを使わず段落で表現）"""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style   = "Table Grid"
    tbl.autofit = False
    left  = tbl.rows[0].cells[0]
    right = tbl.rows[0].cells[1]
    _set_cell_w(left,  label_w)
    _set_cell_w(right, value_w)

    # 左セル（ラベル）
    left.text = ""
    pl = left.paragraphs[0]
    pl.paragraph_format.space_before = Pt(1)
    pl.paragraph_format.space_after  = Pt(1)
    _add_run(pl, label, size=F, bold=True)

    # 右セル（値）
    right.text = ""
    pr = right.paragraphs[0]
    pr.paragraph_format.space_before = Pt(1)
    pr.paragraph_format.space_after  = Pt(1)
    _add_run(pr, value, size=F)

    _remove_borders(tbl)
    return tbl


def _bullets(doc, items, indent_cm=0.3):
    for item in items:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_before     = Pt(0)
        bp.paragraph_format.space_after      = Pt(2)
        bp.paragraph_format.left_indent      = Cm(indent_cm)
        _add_run(bp, item, size=F)


# ── 外務省ページのスクレイピング ──────────────────────────────────────────────
def _abs_url(src: str, base: str) -> str:
    """相対 URL を絶対 URL に変換"""
    if src.startswith("http"):
        return src
    if src.startswith("//"):
        return "https:" + src
    if src.startswith("/"):
        return MOFA_BASE + src
    return base.rsplit("/", 1)[0] + "/" + src


def _fetch_html(url: str) -> str:
    """外務省ページを取得して UTF-8 文字列で返す（gzip 対応）"""
    import gzip as _gzip
    req = urllib.request.Request(url, headers=HEADERS)
    with urllib.request.urlopen(req, timeout=15) as resp:
        raw = resp.read()
    # gzip / deflate を自動展開
    try:
        return _gzip.decompress(raw).decode("utf-8", errors="replace")
    except Exception:
        return raw.decode("utf-8", errors="replace")


def scrape_mofa_images(mofa_url: str) -> tuple:
    """
    外務省の国ページから国旗・周辺地図の画像URLを取得する。

    判定方法:
      - alt 属性に「国旗」が含まれる img → 国旗
      - alt 属性に「地図」が含まれる img → 地域地図
      （外務省ページは全国共通でこの命名規則を採用）

    Returns: (flag_url, map_url)  ※見つからない場合は None
    """
    if not mofa_url or not mofa_url.startswith("http"):
        return None, None

    try:
        html = _fetch_html(mofa_url)
    except Exception as e:
        print(f"  [MOFA] ページ取得失敗: {e}")
        return None, None

    flag_url = None
    map_url  = None

    # imgタグ全体を取得（src と alt をペアで処理）
    for tag in re.finditer(r'<img[^>]+>', html, re.IGNORECASE):
        tag_str = tag.group()
        src_m = re.search(r'\bsrc=["\']([^"\']+)["\']', tag_str, re.I)
        alt_m = re.search(r'\balt=["\']([^"\']*)["\']', tag_str, re.I)
        if not src_m:
            continue
        src = src_m.group(1)
        alt = alt_m.group(1) if alt_m else ""
        full_url = _abs_url(src, mofa_url)

        if not flag_url and "国旗" in alt:
            flag_url = full_url
            print(f"  [MOFA] 国旗画像を検出: {full_url}  (alt={alt})")
        if not map_url and "地図" in alt:
            map_url = full_url
            print(f"  [MOFA] 地図画像を検出: {full_url}  (alt={alt})")
        if flag_url and map_url:
            break

    if not flag_url:
        print("  [MOFA] 国旗画像が見つかりませんでした。")
    if not map_url:
        print("  [MOFA] 地図画像が見つかりませんでした。")

    return flag_url, map_url


def scrape_mofa_ambassadors(mofa_page_url: str) -> dict:
    """
    外務省の国データページ（data.html）の「外交使節」セクションから
    大使名をスクレイピングして返す。

    Returns:
        {
          "japan_ambassador":             "河野章　駐ポーランド日本国特命全権大使",
          "country_ambassador_to_japan":  "パヴェウ・ミレフスキ　駐日ポーランド共和国特命全権大使"
        }
    """
    if not mofa_page_url:
        return {}

    # index.html → data.html に変換
    data_url = re.sub(r'index\.html$', 'data.html', mofa_page_url)
    if data_url == mofa_page_url:
        # index.html が末尾にない場合は /data.html を付加
        data_url = mofa_page_url.rstrip("/") + "/data.html"

    print(f"  [MOFA大使] データページを取得中: {data_url}")
    try:
        html = _fetch_html(data_url)
    except Exception as e:
        print(f"  [MOFA大使] データページ取得失敗: {e}")
        return {}

    # 「外交使節」見出しより後を対象範囲にする
    idx = html.find("外交使節")
    if idx == -1:
        print("  [MOFA大使] 外交使節セクションが見つかりませんでした。")
        return {}

    section = html[idx: idx + 800]

    # <li>…</li> を全件取得（class="none" 等の属性付きタグにも対応）
    items = re.findall(r'<li[^>]*>(.*?)</li>', section, re.DOTALL)

    result = {}
    for item in items:
        # HTMLタグ除去・空白整理
        text = re.sub(r'<[^>]+>', '', item).strip()
        # 先頭の「（1）」「（2）」等を除去
        text = re.sub(r'^[（(]\d+[）)]\s*', '', text).strip()
        if not text:
            continue

        if "駐日" in text:
            result["country_ambassador_to_japan"] = text
        else:
            result["japan_ambassador"] = text

    if result:
        for k, v in result.items():
            print(f"  [MOFA大使] {k}: {v}")
    else:
        print("  [MOFA大使] 大使情報を抽出できませんでした。")

    return result


def download_image(url: str, save_path: Path) -> bool:
    """URL から画像をダウンロードして保存。成功したら True"""
    if not url:
        return False
    try:
        req = urllib.request.Request(url, headers=IMG_HEADERS)
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = resp.read()
        if PILImage:
            img = PILImage.open(BytesIO(data)).convert("RGBA").convert("RGB")
            img.save(str(save_path), "JPEG", quality=92)
        else:
            save_path.write_bytes(data)
        print(f"  [画像] 保存: {save_path.name}")
        return True
    except Exception as e:
        print(f"  [画像] 取得失敗 ({url[:60]}): {e}")
        return False


# ── Claude API ────────────────────────────────────────────────────────────────
def _call_claude(client, prompt: str, max_tokens: int = 2000) -> dict:
    """Claude を呼んで JSON を返す。失敗時は空辞書。"""
    msg = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=max_tokens,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = msg.content[0].text.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s*```\s*$", "", raw)
    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"  [Claude] JSON パースエラー: {e}")
        print(f"  先頭300字: {raw[:300]}")
        return {}


def fetch_basic_data(country: str, client) -> dict:
    print("  [基礎データ] 外務省情報を取得中 (Claude)...")
    return _call_claude(client,
                        BASIC_DATA_PROMPT.format(country=country),
                        max_tokens=1500)


def fetch_economic_data(country: str, basic: dict, client) -> dict:
    print("  [経済データ] 取得中 (Claude)...")
    return _call_claude(
        client,
        ECONOMIC_DATA_PROMPT.format(
            country=country,
            country_en=basic.get("country_name_en", country),
        ),
        max_tokens=1500,
    )


def fetch_japan_relations(country: str, basic: dict, client) -> dict:
    print("  [日本との関係] 取得中 (Claude)...")
    current_year = datetime.now().year
    recent_years = str(current_year - 3)
    return _call_claude(
        client,
        JAPAN_RELATIONS_PROMPT.format(
            country=country,
            country_en=basic.get("country_name_en", country),
            recent_years=recent_years,
        ),
        max_tokens=2000,
    )


# ── Word 文書生成 ──────────────────────────────────────────────────────────────
def generate_doc(country_input: str,
                 basic: dict,
                 eco: dict,
                 jpn: dict,
                 flag_path,
                 map_path,
                 output_path: Path):

    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(1.0)

    country_ja = basic.get("country_name_ja", country_input)
    date_str   = datetime.now().strftime("%Y年%m月%d日")

    # ── 添付番号 ──────────────────────────────────────────────────────────────
    p_attach = _para(doc, "添付３", size=14, bold=True,
                     align=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)
    bdr = OxmlElement("w:bdr")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), "4")
    bdr.set(qn("w:space"), "0")
    bdr.set(qn("w:color"), "000000")
    p_attach.runs[0]._r.get_or_add_rPr().append(bdr)

    # ── ① タイトル（左：国名＋日付、右：国旗） ──────────────────────────────
    tbl_title = doc.add_table(rows=1, cols=2)
    tbl_title.style   = "Table Grid"
    tbl_title.autofit = False
    tc_left  = tbl_title.rows[0].cells[0]
    tc_right = tbl_title.rows[0].cells[1]
    _set_cell_w(tc_left,  13.5)
    _set_cell_w(tc_right,  3.5)
    _remove_borders(tbl_title)

    # 左セル：タイトル＋日付
    tc_left.text = ""
    tp = tc_left.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after  = Pt(2)
    _add_run(tp, f"{country_ja}　概　況", size=18, bold=True)

    # 右セル：国旗
    tc_right.text = ""
    fp = tc_right.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after  = Pt(0)
    if flag_path and Path(flag_path).exists():
        try:
            fp.add_run().add_picture(str(flag_path), width=Cm(3.0))
        except Exception as e:
            print(f"  [Word] タイトル行国旗貼付エラー: {e}")

    # ── ② 基礎データ（ラベル列 ＋ 値列）＋ 地図（右列・全行スパン） ────────────
    _section_title(doc, "基　礎　デ　ー　タ")

    def _list_val(lst, sep="、"):
        if not lst:
            return "不明"
        items = [str(x) for x in lst if x and x != "不明"]
        return sep.join(items) if items else "不明"

    basic_rows = [
        ("正 式 国 名", basic.get("country_name_ja", "不明")),
        ("英　語　名",  basic.get("country_name_en", "不明")),
        ("面　　　積",  basic.get("area_km2",         "不明")),
        ("人　　　口",  basic.get("population",       "不明")),
        ("首　　　都",  basic.get("capital",          "不明")),
        ("民　　　族",  _list_val(basic.get("ethnic_groups", []))),
        ("言　　　語",  _list_val(basic.get("languages",    []))),
        ("宗　　　教",  _list_val(basic.get("religions",    []))),
    ]

    # 3列テーブル: [ラベル(3.0cm)] [値(8.3cm)] [地図(5.5cm・全行スパン)]
    # → 値が長くても「値列」内で折り返し、ラベル列に食み出さない
    NUM_ROWS = len(basic_rows)
    LABEL_W  = 3.0
    VALUE_W  = 8.5
    MAP_W    = 6.5

    tbl_basic = doc.add_table(rows=NUM_ROWS, cols=3)
    tbl_basic.style   = "Table Grid"
    tbl_basic.autofit = False
    _remove_borders(tbl_basic)

    for row in tbl_basic.rows:
        _set_cell_w(row.cells[0], LABEL_W)
        _set_cell_w(row.cells[1], VALUE_W)
        _set_cell_w(row.cells[2], MAP_W)

    # 右列（地図）を全行マージして1セルに
    map_cell = tbl_basic.rows[0].cells[2]
    for i in range(1, NUM_ROWS):
        map_cell = map_cell.merge(tbl_basic.rows[i].cells[2])

    map_cell.text = ""
    mp = map_cell.paragraphs[0]
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_before = Pt(8)
    mp.paragraph_format.space_after  = Pt(0)
    if map_path and Path(map_path).exists():
        try:
            mp.add_run().add_picture(str(map_path), width=Cm(5.8))
        except Exception as e:
            print(f"  [Word] 地図貼付エラー: {e}")
    else:
        _add_run(mp, "（地図取得不可）", size=F - 1,
                 color=RGBColor(0x80, 0x80, 0x80))

    # データ行：ラベル列（太字）＋ 値列（長い値も列内で折り返して揃う）
    for i, (label, value) in enumerate(basic_rows):
        lc = tbl_basic.rows[i].cells[0]
        vc = tbl_basic.rows[i].cells[1]

        lc.text = ""
        pl = lc.paragraphs[0]
        pl.paragraph_format.space_before = Pt(2)
        pl.paragraph_format.space_after  = Pt(2)
        _add_run(pl, f"{label}：", bold=True)

        vc.text = ""
        pv = vc.paragraphs[0]
        pv.paragraph_format.space_before = Pt(2)
        pv.paragraph_format.space_after  = Pt(2)
        _add_run(pv, value)

    # ── ③ 経済状況 ────────────────────────────────────────────────────────────
    _section_title(doc, "経　済　状　況")

    def _eco_row(label, val):
        _kv_row(doc, label, val or "不明", label_w=4.5, value_w=11.5)

    _eco_row("主 要 産 業",
             _list_val(eco.get("major_industries", []), "、"))
    _eco_row("GDP（名目）",          eco.get("gdp",              "不明"))
    _eco_row("一人当たり GDP",        eco.get("gdp_per_capita",   "不明"))
    _eco_row("経 済 成 長 率",       eco.get("gdp_growth_rate",  "不明"))
    _eco_row("物 価 上 昇 率",       eco.get("inflation_rate",   "不明"))
    _eco_row("失　業　率",            eco.get("unemployment_rate","不明"))
    _eco_row("総輸出額",              eco.get("total_exports",    "不明"))
    _eco_row("総輸入額",              eco.get("total_imports",    "不明"))
    _eco_row("主要輸出品目",
             _list_val(eco.get("major_export_items", []), "、"))
    _eco_row("主要輸入品目",
             _list_val(eco.get("major_import_items", []), "、"))
    _eco_row("主要貿易相手国",
             _list_val(eco.get("major_trade_partners", []), "、"))



    # ── ④ 日本との関係 ───────────────────────────────────────────────────────
    _section_title(doc, "日　本　と　の　関　係")

    # 要人往来（時系列順にソート）
    def _visit_sort_key(v):
        """date 文字列（例: "2023年3月"）を (year, month) の数値タプルに変換"""
        d = v.get("date", "")
        y = re.search(r'(\d{4})年', d)
        m = re.search(r'(\d{1,2})月', d)
        return (int(y.group(1)) if y else 0, int(m.group(1)) if m else 0)

    vip_visits = sorted(
        [v for v in jpn.get("vip_visits", [])
         if isinstance(v, dict) and v.get("summary", "不明") != "不明"],
        key=_visit_sort_key,
        reverse=True   # 最新を冒頭に
    )
    if vip_visits:
        _para(doc, "◆ 要人往来（直近2〜3年）",
              size=F, bold=True, space_before=4, space_after=2)
        for v in vip_visits:
            date  = v.get("date", "")
            summ  = v.get("summary", "")
            _para(doc, f"・{date}　{summ}", size=F,
                  space_before=0, space_after=2)

    # 条約・協定（直近2〜3年のもののみ。該当なしの場合はセクション自体を省略）
    cutoff_year = datetime.now().year - 3
    treaties = [
        t for t in jpn.get("treaties", [])
        if isinstance(t, dict)
        and t.get("name", "不明") != "不明"
        and re.search(r'(\d{4})', t.get("year", ""))
        and int(re.search(r'(\d{4})', t.get("year", "")).group(1)) >= cutoff_year
    ]
    if treaties:
        _para(doc, "◆ 条約・協定（直近2〜3年）",
              size=F, bold=True, space_before=4, space_after=2)
        for t in treaties:
            _para(doc, f"・{t.get('name','')}（{t.get('year','')}）",
                  size=F, space_before=0, space_after=2)

    # 日本との貿易
    _para(doc, "◆ 日本との貿易",
          size=F, bold=True, space_before=4, space_after=2)
    _kv_row(doc, "日本からの輸出額",
            jpn.get("japan_exports_to",   "不明"), label_w=5.5, value_w=10.5)
    _kv_row(doc, "日本への輸入額",
            jpn.get("japan_imports_from", "不明"), label_w=5.5, value_w=10.5)
    _kv_row(doc, "日本からの輸出品目",
            _list_val(jpn.get("major_japan_export_items", []), "、"),
            label_w=5.5, value_w=10.5)
    _kv_row(doc, "日本への輸入品目",
            _list_val(jpn.get("major_japan_import_items", []), "、"),
            label_w=5.5, value_w=10.5)

    # 日系企業
    _para(doc, "◆ 日系企業",
          size=F, bold=True, space_before=4, space_after=2)
    _kv_row(doc, "進 出 企 業 数",
            jpn.get("japanese_companies_count", "不明"),
            label_w=5.5, value_w=10.5)
    notable = [c for c in jpn.get("japanese_companies_notable", [])
               if c and c != "不明"]
    if notable:
        _kv_row(doc, "主な進出企業",
                "　".join(notable),
                label_w=5.5, value_w=10.5)

    # 在留邦人
    _para(doc, "◆ 在留邦人",
          size=F, bold=True, space_before=4, space_after=2)
    _kv_row(doc, "在 留 邦 人 数",
            jpn.get("japanese_residents", "不明"),
            label_w=5.5, value_w=10.5)

    # 大使
    _para(doc, "◆ 大使",
          size=F, bold=True, space_before=4, space_after=2)
    _kv_row(doc, f"駐{country_ja}日本大使",
            jpn.get("japan_ambassador", "不明"),
            label_w=5.5, value_w=10.5)
    _kv_row(doc, f"駐日{country_ja}大使",
            jpn.get("country_ambassador_to_japan", "不明"),
            label_w=5.5, value_w=10.5)

    # ── フッター ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    _para(doc, "以　上", size=F, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)

    doc.save(str(output_path))


# ── メイン ────────────────────────────────────────────────────────────────────
def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)

    print("=" * 55)
    print("  国概況資料ジェネレーター")
    print("=" * 55)
    try:
        country = input(
            "\n対象国を入力してください。\n"
            "（例: アメリカ、中国、フランス、インド、サウジアラビア）\n> "
        ).strip()
    except (EOFError, KeyboardInterrupt):
        print("\nキャンセルされました。")
        return
    if not country:
        sys.exit("国名が入力されませんでした。")

    client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)

    # 1. 基礎データ取得（外務省情報ベース）
    basic = fetch_basic_data(country, client)
    if not basic:
        sys.exit("基礎データの取得に失敗しました。")

    country_ja = basic.get("country_name_ja", country)
    print(f"\n  対象国: {country_ja}（{basic.get('country_name_en','')}）")
    print(f"  外務省URL: {basic.get('mofa_page_url','不明')}")

    # 2. 外務省ページから国旗・地図画像を取得
    safe_name = re.sub(r'[\\/:*?"<>|\s]', "_", country_ja)
    flag_path = IMG_DIR / f"{safe_name}_flag.jpg"
    map_path  = IMG_DIR / f"{safe_name}_map.jpg"

    print("\n  [画像] 外務省ページをスクレイピング中...")
    mofa_url = basic.get("mofa_page_url", "")
    flag_url_scraped, map_url_scraped = scrape_mofa_images(mofa_url)

    # 国旗: スクレイピング結果 → Claude提示URL → Wikipedia URL の順に試行
    flag_ok = False
    for url in filter(None, [
        flag_url_scraped,
        basic.get("mofa_flag_img_url"),
        basic.get("wikipedia_flag_url"),
    ]):
        if download_image(url, flag_path):
            flag_ok = True
            break
    if not flag_ok:
        print("  [画像] 国旗を取得できませんでした。")
        flag_path = None

    # 地図: スクレイピング結果 → Claude提示URL の順に試行
    map_ok = False
    for url in filter(None, [
        map_url_scraped,
        basic.get("mofa_map_img_url"),
    ]):
        if download_image(url, map_path):
            map_ok = True
            break
    if not map_ok:
        print("  [画像] 周辺地図を取得できませんでした。")
        map_path = None

    # 3. 経済データ取得
    eco = fetch_economic_data(country, basic, client)

    # 4. 日本との関係取得
    jpn = fetch_japan_relations(country, basic, client)

    # 4b. 外務省 data.html から大使名をスクレイピングして上書き（優先）
    print("\n  [大使情報] 外務省データページをスクレイピング中...")
    ambassador_data = scrape_mofa_ambassadors(basic.get("mofa_page_url", ""))
    if ambassador_data:
        jpn.update(ambassador_data)   # スクレイピング結果で Claude の推測を上書き

    # 5. Word 文書生成
    date_tag    = datetime.now().strftime("%Y%m%d")
    output_path = OUT_DIR / f"{safe_name}_概況_{date_tag}.docx"

    print("\n  [Word] 文書を生成中...")
    try:
        generate_doc(country, basic, eco, jpn, flag_path, map_path, output_path)
        print(f"\n保存完了: {output_path}")
    except Exception as e:
        sys.exit(f"文書生成に失敗しました: {e}")


if __name__ == "__main__":
    main()
