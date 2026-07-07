#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
面談資料インタラクティブ作成プログラム
PDF の黒塗り箇所を入力し、Word (.docx) 文書を生成します。
"""

import os
import sys
import json

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("エラー: python-docx が必要です。")
    print("  pip install python-docx")
    sys.exit(1)


# ═══════════════════════════════════════════
#  入力ヘルパー
# ═══════════════════════════════════════════

def ask(label, required=True):
    while True:
        val = input(f"  {label}: ").strip()
        if val or not required:
            return val
        print("  ※ 入力してください。")


def ask_lines(label, required=False):
    """空行が入力されるまで複数行を受け付ける"""
    print(f"  {label}（空行で完了）:")
    lines = []
    while True:
        line = input("    > ").strip()
        if not line:
            if required and not lines:
                print("    ※ 最低1行入力してください。")
                continue
            return lines
        lines.append(line)


# ═══════════════════════════════════════════
#  Word ヘルパー
# ═══════════════════════════════════════════

JP_FONT = "MS 明朝"
EN_FONT = "Times New Roman"


def set_font(run, size=10.5, bold=False, font_jp=None):
    jp = font_jp if font_jp else JP_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = EN_FONT
    rpr = run._r.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    rf.set(qn("w:eastAsia"), jp)
    rf.set(qn("w:ascii"),    EN_FONT)
    rf.set(qn("w:hAnsi"),    EN_FONT)


def write_cell(cell, text, size=10.5, bold=False,
               align=WD_ALIGN_PARAGRAPH.LEFT, font_jp=None):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size, bold, font_jp=font_jp)


def write_cell_lines(cell, lines, size=10.5, font_jp=None):
    """複数行をセルに書き込む（1行目は既存の段落を利用）。行間を詰める。"""
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        set_font(run, size, font_jp=font_jp)
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        pf.line_spacing = Pt(size + 7)


def _get_or_add_tblPr(tbl):
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    return tblPr


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = _get_or_add_tblPr(tbl)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tb = OxmlElement("w:tblBorders")
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{e}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:color"), "FFFFFF")
        tb.append(el)
    tblPr.append(tb)
    # セルレベルの罫線も白に設定
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            old_cb = tcPr.find(qn("w:tcBorders"))
            if old_cb is not None:
                tcPr.remove(old_cb)
            cb = OxmlElement("w:tcBorders")
            for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{e}")
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "0")
                el.set(qn("w:color"), "FFFFFF")
                cb.append(el)
            tcPr.append(cb)


def add_box_borders(table):
    """テーブルの全罫線を白（非表示）にする"""
    tbl = table._tbl
    tblPr = _get_or_add_tblPr(tbl)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tb = OxmlElement("w:tblBorders")
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{e}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:color"), "FFFFFF")
        tb.append(el)
    tblPr.append(tb)


def set_table_indent(table, indent_cm):
    """テーブルに左インデントを設定する"""
    tbl = table._tbl
    tblPr = _get_or_add_tblPr(tbl)
    old = tblPr.find(qn("w:tblInd"))
    if old is not None:
        tblPr.remove(old)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), str(int(indent_cm * 566.93)))
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)


def calc_gaiyou_height_cm(d):
    """概要テーブルの高さを推定する（テキスト折り返しを考慮）"""
    # 右欄幅: TEXT_W - TBL_INDENT - LBL_W - セル左右パディング(0.19cm×2)
    CONT_W_CM    = 17.0 - 0.5 - 4.0 - 0.38   # ≈ 12.12cm
    FULL_W_CM    = 0.423   # 游明朝12pt 全角1文字
    HALF_W_CM    = 0.212   # 游明朝12pt 半角1文字

    def text_lines(text):
        """折り返しを考慮した視覚的行数を返す"""
        if not text:
            return 1
        w = 0
        lines = 1
        for ch in str(text):
            cw = HALF_W_CM if ch.isascii() else FULL_W_CM
            w += cw
            if w > CONT_W_CM:
                lines += 1
                w = cw
        return lines

    def content_lines(value):
        """セルコンテンツ（文字列またはリスト）の合計行数"""
        if isinstance(value, list):
            return sum(text_lines(s) for s in value) if value else 1
        return text_lines(value)

    dt = (f"{d['mtg_month']}月{d['mtg_day']}日({d['mtg_dow']}) "
          f"{d['mtg_start']}－{d['mtg_end']} （於：{d['mtg_venue']}）")

    purposes = [f"① {d['purpose1']}"]
    if d.get('purpose2'):
        purposes.append(f"② {d['purpose2']}")

    rows = [
        dt,
        d.get('attendees') or [''],
        purposes,
        d.get('language', ''),
        d.get('flow', ''),
        d.get('dress_code', ''),
        d.get('reference') or [''],
    ]

    LINE_SP_PT  = 19    # 行間（size+7 と統一）
    SPACE_AFT   = 12    # space_after
    PT_TO_CM    = 0.03528
    # posOffset=-80000EMU 分（上にずらしている量）を高さに加算
    POS_OFFSET_CM = 120000 / 360000  # 0.333cm
    BUFFER_CM   = 0.3

    total_pt = sum(content_lines(r) * LINE_SP_PT + SPACE_AFT for r in rows)
    return total_pt * PT_TO_CM + POS_OFFSET_CM + BUFFER_CM


def insert_rect_shape(para, width_cm, height_cm, h_offset_cm=0, shape_id=1):
    """段落に黒枠・透明塗りの矩形図形（DrawingML）を挿入する"""
    from lxml import etree
    CM = 360000
    w = int(width_cm * CM)
    h = int(height_cm * CM)
    x = int(h_offset_cm * CM)
    xml_str = (
        f'<w:r'
        f' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
        f' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
        f' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        f'<w:drawing>'
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0"'
        f' simplePos="0" relativeHeight="251659264"'
        f' behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="column"><wp:posOffset>{x}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="paragraph"><wp:posOffset>-120000</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{w}" cy="{h}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="{shape_id}" name="Rectangle {shape_id}"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic>'
        f'<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        f'<wps:wsp>'
        f'<wps:cNvSpPr><a:spLocks noChangeArrowheads="1"/></wps:cNvSpPr>'
        f'<wps:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{w}" cy="{h}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'<a:noFill/>'
        f'<a:ln w="6350"><a:solidFill><a:srgbClr val="000000"/></a:solidFill></a:ln>'
        f'</wps:spPr>'
        f'<wps:bodyPr/>'
        f'</wps:wsp>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:anchor>'
        f'</w:drawing>'
        f'</w:r>'
    )
    r_elem = etree.fromstring(xml_str)
    para._p.append(r_elem)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


# ═══════════════════════════════════════════
#  文書生成
# ═══════════════════════════════════════════

def create_document(d, out_path):
    doc = Document()

    # A4、余白設定
    sec = doc.sections[0]
    sec.page_width    = Cm(21.0)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)

    TEXT_W = 17.0  # テキスト幅 cm (21 - 2.0*2)

    # ── ① ヘッダー：宛先（行0）/ 年月日・差出人（行1・2） ────────
    # 行0: 宛先のみ
    # 行1: 右側に年月日
    # 行2: 右側に差出人
    ht = doc.add_table(rows=3, cols=2)
    ht.autofit = False
    remove_table_borders(ht)
    set_col_widths(ht, [TEXT_W * 0.5, TEXT_W * 0.5])

    # 行0：宛先（行高を自動にして文字が全部見えるよう設定）
    write_cell(ht.cell(0, 0), f"{d['recipient']}　殿", font_jp="游明朝", size=14)
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    ht.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AUTO

    # 行1：年月日（右）
    write_cell(ht.cell(1, 1),
               f"{d['year']}年{d['month']}月{d['day']}日",
               align=WD_ALIGN_PARAGRAPH.RIGHT, font_jp="游明朝", size=11.5)

    # 行2：差出人（右）
    write_cell(ht.cell(2, 1), d['sender'],
               align=WD_ALIGN_PARAGRAPH.RIGHT, font_jp="游明朝", size=11.5)

    # 年月日・差出人の間隔設定
    for cell in ht.rows[1].cells:        # 年月日行：下に少し余白
        for para in cell.paragraphs:
            pf = para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after  = Pt(5)
            pf.line_spacing = Pt(13)
    for cell in ht.rows[2].cells:        # 差出人行
        for para in cell.paragraphs:
            pf = para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after  = Pt(0)
            pf.line_spacing = Pt(13)

    # 宛先行（行0）の段落余白もリセット
    for para in ht.cell(0, 0).paragraphs:
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(2)

    # ── ② 件名（中央寄せ・下線・太字） ───────────────────────
    doc.add_paragraph()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(f"{d['meeting_partner']}との面談について")
    r.underline = True
    set_font(r, size=14, bold=True, font_jp="游明朝")

    # ── ③ 本文（固定） ────────────────────────────────────
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = Pt(16)
    set_font(p.add_run("本件、ご対応くださり、誠にありがとうございます。"),
             size=12, font_jp="游明朝")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = Pt(16)
    set_font(p.add_run(
        "概要と懇談のトピックスについてご相談させて頂きたく、よろしくお願いいたします。"),
             size=12, font_jp="游明朝")

    # ── ④ <概要> セクション ──────────────────────────────
    doc.add_paragraph()
    ph = doc.add_paragraph()
    ph.paragraph_format.space_after = Pt(0)
    set_font(ph.add_run("＜概要＞"), size=13, bold=True, font_jp="游明朝")

    # ＜概要＞と表の間のスペーサー（約0.5cm）
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after  = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(8.5)
    spacer.add_run().font.size = Pt(1)

    # 矩形図形用アンカー段落（高さゼロ・空行にならないよう設定）
    rect_para = doc.add_paragraph()
    rect_para.paragraph_format.space_before = Pt(0)
    rect_para.paragraph_format.space_after  = Pt(0)
    rect_para.paragraph_format.line_spacing = Pt(1)
    _r = rect_para.add_run()
    _r.font.size = Pt(1)
    TBL_INDENT = 0.5  # 左インデント cm
    # 矩形の高さをデータから推定し、底辺が表の底辺と合うよう設定
    rect_height = calc_gaiyou_height_cm(d)
    insert_rect_shape(rect_para, width_cm=TEXT_W,
                      height_cm=rect_height, h_offset_cm=0, shape_id=1)

    # 2列テーブル（ラベル列 + 内容列）、全罫線白
    LBL_W  = 3.577
    CONT_W = TEXT_W - TBL_INDENT - LBL_W  # 右枠固定のため幅を縮める

    gt = doc.add_table(rows=7, cols=2)
    gt.autofit = False
    add_box_borders(gt)
    set_col_widths(gt, [LBL_W, CONT_W])
    set_table_indent(gt, TBL_INDENT)

    def fill_row(row_idx, label, value_or_lines):
        write_cell(gt.cell(row_idx, 0), label, size=12, bold=False, font_jp="游明朝")
        cell = gt.cell(row_idx, 1)
        if isinstance(value_or_lines, list):
            write_cell_lines(cell, value_or_lines if value_or_lines else [""],
                             size=12, font_jp="游明朝")
        else:
            write_cell(cell, value_or_lines, size=12, font_jp="游明朝")
        # 全セルの段落に明示的な行間を設定（高さ計算と一致させる）
        for c in (gt.cell(row_idx, 0), gt.cell(row_idx, 1)):
            for para in c.paragraphs:
                para.paragraph_format.line_spacing = Pt(19)
            c.paragraphs[-1].paragraph_format.space_after = Pt(12)

    # 日時・場所
    dt = (f"{d['mtg_month']}月{d['mtg_day']}日({d['mtg_dow']})　"
          f"{d['mtg_start']}－{d['mtg_end']}　（於：{d['mtg_venue']}）")

    # 目的
    purposes = [f"① {d['purpose1']}"]
    if d['purpose2']:
        purposes.append(f"② {d['purpose2']}")

    fill_row(0, "日時・場所　：", dt)
    fill_row(1, "出席者　　　：", d['attendees'])
    fill_row(2, "目的　　　　：", purposes)
    # ①② の後ろで折り返しを揃えるぶら下げインデント
    _hang = Cm(0.635)
    for _p in gt.cell(2, 1).paragraphs:
        _p.paragraph_format.left_indent = _hang
        _p.paragraph_format.first_line_indent = -_hang
    fill_row(3, "言語　　　　：", d['language'])
    fill_row(4, "流れ　　　　：", d['flow'])
    fill_row(5, "服装　　　　：", d['dress_code'])
    fill_row(6, "（ご参考）",      d['reference'])

    # ── ⑤ <添付資料> セクション ──────────────────────────
    doc.add_paragraph()
    pa = doc.add_paragraph()
    pa.paragraph_format.space_before = Pt(0)
    pa.paragraph_format.space_after  = Pt(0)
    pa.paragraph_format.line_spacing = Pt(19)
    set_font(pa.add_run("＜添付資料＞"), size=10.5, bold=True, font_jp="游明朝")

    NUMS = ["1.", "2.", "3."]
    att_parts = [f"{NUMS[i]} {att}" for i, att in enumerate(d['attachments'][:3]) if att]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = Pt(19)
    p.paragraph_format.left_indent  = Cm(0.423)
    set_font(p.add_run("　".join(att_parts)), size=10.5, font_jp="游明朝")

    # ── ⑥ 以上 ────────────────────────────────────────
    ep = doc.add_paragraph()
    ep.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(ep.add_run("以　　上"), size=12, font_jp="游明朝")

    # ── 保存 ───────────────────────────────────────────
    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    while True:
        try:
            doc.save(out_path)
            print(f"\n[完了] 保存しました: {out_path}\n")
            break
        except PermissionError:
            print(f"\n[警告] 保存できません: {out_path}")
            print("   Word でファイルが開いている場合は閉じてください。")
            ans = input("   閉じたら Enter を押す（別名保存は b を入力）: ").strip().lower()
            if ans == "b":
                base, ext = os.path.splitext(out_path)
                from datetime import datetime
                out_path = f"{base}_{datetime.now().strftime('%H%M%S')}{ext}"
                print(f"   別名: {out_path}")
            # そのままEnterならリトライ


# ═══════════════════════════════════════════
#  入力データ 保存・読み込み
# ═══════════════════════════════════════════

SAVE_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "前回の入力データ.json")


def save_data(d):
    with open(SAVE_FILE, "w", encoding="utf-8") as f:
        json.dump(d, f, ensure_ascii=False, indent=2)


def load_data():
    if not os.path.exists(SAVE_FILE):
        return None
    with open(SAVE_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def show_summary(d):
    print()
    print("  ┌─────────────────────────────────────────────")
    print(f"  │ 宛先       : {d['recipient']}　殿")
    print(f"  │ 日付       : {d['year']}年{d['month']}月{d['day']}日")
    print(f"  │ 差出人     : {d['sender']}")
    print(f"  │ 件名       : {d['meeting_partner']}との面談について")
    print(f"  │ 日時・場所 : {d['mtg_month']}月{d['mtg_day']}日({d['mtg_dow']}) "
          f"{d['mtg_start']}－{d['mtg_end']} 於:{d['mtg_venue']}")
    print(f"  │ 出席者     : {' / '.join(d['attendees'])}")
    print(f"  │ 目的①  : {d['purpose1']}")
    if d.get('purpose2'):
        print(f"  │ 目的②  : {d['purpose2']}")
    print(f"  │ 言語       : {d['language']}")
    print(f"  │ 流れ       : {d['flow']}")
    print(f"  │ 服装       : {d['dress_code']}")
    if d.get('reference'):
        print(f"  │ ご参考     : {' / '.join(d['reference'])}")
    if d.get('attachments'):
        print(f"  │ 添付資料   : {' / '.join(d['attachments'])}")
    print("  └─────────────────────────────────────────────")
    print()


def collect_input():
    d = {}

    print("【基本情報】")
    d["recipient"]       = ask("宛先（〇〇殿 の 〇〇）")
    d["year"]            = ask("年（例：2026）")
    d["month"]           = ask("月（例：5）")
    d["day"]             = ask("日（例：21）")
    d["sender"]          = ask("差出人（氏名・所属など）")

    print("\n【件名】")
    d["meeting_partner"] = ask("面談相手（〇〇との面談 の 〇〇）")

    print("\n【概要：日時・場所】")
    d["mtg_month"] = ask("面談の月（例：5）")
    d["mtg_day"]   = ask("面談の日（例：30）")
    d["mtg_dow"]   = ask("曜日（例：土）")
    d["mtg_start"] = ask("開始時刻（例：14:00）")
    d["mtg_end"]   = ask("終了時刻（例：15:30）")
    d["mtg_venue"] = ask("場所（例：〇〇ホテル〇〇ルーム）")

    print("\n【概要：出席者】")
    d["attendees"] = ask_lines("出席者（1行ずつ入力）", required=True)

    print("\n【概要：目的】")
    d["purpose1"] = ask("目的①")
    d["purpose2"] = ask("目的②（不要なら空欄）", required=False)

    print("\n【概要：言語・流れ・服装】")
    d["language"]   = ask("言語（例：日本語）")
    d["flow"]       = ask("流れ（例：先方挨拶→概要説明→質疑）")
    d["dress_code"] = ask("服装（例：スーツ）")

    print("\n【概要：ご参考】（任意）")
    d["reference"] = ask_lines("ご参考（1行ずつ入力）", required=False)

    print("\n【添付資料】（不要な項目は空欄）")
    d["attachments"] = []
    for i in range(1, 4):
        a = ask(f"添付資料{i}", required=False)
        if a:
            d["attachments"].append(a)

    return d


# ═══════════════════════════════════════════
#  メイン
# ═══════════════════════════════════════════

def main():
    print()
    print("=" * 58)
    print("    面談資料 インタラクティブ作成プログラム")
    print("=" * 58)

    # 前回データの確認
    saved = load_data()
    d = None
    if saved:
        print("\n  前回の入力データが見つかりました。")
        show_summary(saved)
        ans = input("  このデータで作成しますか？ [y = はい / n = 新規入力]: ").strip().lower()
        if ans in ("y", "yes", ""):
            d = saved
        else:
            print()

    if d is None:
        print("  ※ PDF の黒塗り箇所を順番に入力してください。")
        print("  ※ 任意項目は空欄のまま Enter で省略できます。")
        print()
        d = collect_input()

    # 出力先
    default_out = os.path.join(os.path.expanduser("~"), "Desktop", "面談資料.docx")
    print(f"\n出力先（デフォルト）: {default_out}")
    custom = input("  変更する場合はパスを入力（そのままでよければ Enter）: ").strip()
    out_path = custom if custom else default_out

    create_document(d, out_path)

    # 入力データを保存
    save_data(d)
    print(f"  [保存] 入力データを保存しました: {SAVE_FILE}")


if __name__ == "__main__":
    main()
