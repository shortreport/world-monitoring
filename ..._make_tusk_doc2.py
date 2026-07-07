from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from pathlib import Path

OUTPUT = Path(r"C:\Users\shondo\Desktop\トゥスク首相_略歴書.docx")
PHOTO  = Path(r"C:\Users\shondo\Desktop\agent_project\_tusk_photo_fixed.jpg")
F = 12  # 統一フォントサイズ

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "FFFFFF")
        tblBorders.append(b)
    tblPr.append(tblBorders)

doc = Document()

sec = doc.sections[0]
sec.page_width    = Cm(21)
sec.page_height   = Cm(29.7)
sec.top_margin    = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.0)

def para(text="", size=F, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_after=3, underline=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.underline = underline
        run.font.name = "MS明朝"
        _set_east_asian_font(run, "MS明朝")
        if color:
            run.font.color.rgb = color
    return p

def _set_east_asian_font(run, font_name):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)

def set_cell_font(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(F)
    run.font.bold = bold
    run.font.name = "MS明朝"
    _set_east_asian_font(run, "MS明朝")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

def add_bullets(items):
    for item in items:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after  = Pt(3)
        run = bp.add_run(item)
        run.font.size = Pt(F)
        run.font.name = "MS明朝"
        _set_east_asian_font(run, "MS明朝")

def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")

# ── タイトル ─────────────────────────────────────────────────
para("履　　　歴　　　書", size=16, bold=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

# ── 基本情報 + 写真 ──────────────────────────────────────────
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
tbl.autofit = False

left  = tbl.rows[0].cells[0]
right = tbl.rows[0].cells[1]
set_cell_width(left,  12.8)
set_cell_width(right,  4.2)

def add_info_line(cell, label, value, first=False):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(label)
    r1.font.size = Pt(F)
    r1.font.name = "MS明朝"
    _set_east_asian_font(r1, "MS明朝")
    r2 = p.add_run(f"　{value}")
    r2.font.size = Pt(F)
    r2.font.bold = True
    r2.font.name = "MS明朝"
    _set_east_asian_font(r2, "MS明朝")

add_info_line(left, "氏　　名", "ドナルド・トゥスク", first=True)
add_info_line(left, "生年月日", "1957年 4月 22日")
add_info_line(left, "国　　籍", "ポーランド共和国")
add_info_line(left, "現　　職", "ポーランド共和国首相（2023年12月～）")

remove_table_borders(tbl)
rp = right.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
rp.paragraph_format.space_before = Pt(0)
rp.add_run().add_picture(str(PHOTO), width=Cm(4.0))

doc.add_paragraph()

# ── 学歴 ────────────────────────────────────────────────────
para("学　　　　　歴", size=F, bold=True, space_after=2, underline=True)

edu = doc.add_table(rows=1, cols=2)
edu.style = "Table Grid"
edu.autofit = False
set_cell_font(edu.rows[0].cells[0], "1975年 ～ 1980年")
set_cell_font(edu.rows[0].cells[1], "グダンスク大学　歴史学部卒業")
set_cell_width(edu.rows[0].cells[0], 4.7)
set_cell_width(edu.rows[0].cells[1], 11.8)
remove_table_borders(edu)

doc.add_paragraph()

# ── 主な職歴（功績・現職に関係するもののみ）──────────────────
para("主　な　職　歴", size=F, bold=True, space_after=2, underline=True)

careers = [
    ("1980年",   "民主化運動「連帯」に参加"),
    ("1989年",   "自由民主会議（KLD）を共同設立し政界入り"),
    ("2001年",   "「市民プラットフォーム（PO）」を共同設立"),
    ("2007年",   "ポーランド共和国首相に就任（第一期）"),
    ("2011年",   "首相再選（1989年以降初の再選首相）"),
    ("2014年",   "欧州理事会議長に就任（ポーランド人として初）"),
    ("2019年",   "欧州理事会議長　退任"),
    ("2023年",   "ポーランド共和国首相に再就任（現在に至る）"),
]

ctbl = doc.add_table(rows=len(careers), cols=2)
ctbl.style = "Table Grid"
ctbl.autofit = False

for i, (date, desc) in enumerate(careers):
    set_cell_font(ctbl.rows[i].cells[0], date)
    set_cell_font(ctbl.rows[i].cells[1], desc)
    set_cell_width(ctbl.rows[i].cells[0], 4.7)
    set_cell_width(ctbl.rows[i].cells[1], 11.8)
remove_table_borders(ctbl)

doc.add_paragraph()

# ── 関心事項・主要政策 ────────────────────────────────────────
para("関　心　事　項・主　要　政　策", size=F, bold=True, space_after=4, underline=True)

add_bullets([
    "法の支配・司法の独立回復（PiS政権期の違憲的司法機関の解体）",
    "EU統合の深化とポーランド・EU関係の修復（EU資金の凍結解除を実現）",
    "ウクライナへの全面支援とEU加盟プロセスの推進",
    "欧州防衛の強化（NATO加盟国中最大のGDP比4.7%の国防費を計上）",
    "欧州防衛の対米自律性向上（「欧州防衛銀行」構想を提唱）",
    "トランプ政権下での対米関係の安定的維持",
])

doc.add_paragraph()

# ── To Do ────────────────────────────────────────────────────
para("To Do　（歓迎される話題・姿勢）", size=F, bold=True, space_after=4, underline=True)

add_bullets([
    "ウクライナへの支持と連帯を明確に表明する",
    "法の支配・民主主義という共通の価値観を前面に出す",
    "ポーランドの欧州防衛への貢献（GDP比4.7%の国防費）を称賛する",
    "EU統合の強化と欧州の戦略的自律性について積極的に議論する",
    "日欧の経済安全保障・サプライチェーン強靭化での連携を提案する",
    "対中国リスクへの対応において日欧の利害が一致することを示す",
    "欧州理事会議長としての実績に対して敬意を示す",
])

doc.add_paragraph()

# ── Not to Do ────────────────────────────────────────────────
para("Not to Do　（避けるべき言動）", size=F, bold=True, space_after=4, underline=True)

add_bullets([
    "ウクライナに領土的譲歩を求める停戦案に同調するような発言をしない",
    "NATOの結束や米国の欧州へのコミットメントに疑問を呈さない",
    "PiS（法と正義）政権期の政策を肯定的に評価しない",
    "EU統合に懐疑的な印象を与える発言を避ける",
    "欧州の安全保障問題を日本に無関係なテーマとして扱わない",
    "ロシアとの対話・関係改善を優先すべきと示唆しない",
    "移民・難民問題でEUの立場と相容れない強硬論を展開しない",
])

doc.add_paragraph()

# ── 作成日 ───────────────────────────────────────────────────
date_str = datetime.now().strftime("%Y年　%m月　%d日　作成")
para(date_str, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT,
     color=RGBColor(0x60, 0x60, 0x60), space_after=0)

doc.save(str(OUTPUT))
print(f"保存完了: {OUTPUT}")
