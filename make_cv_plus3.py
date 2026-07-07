"""
make_cv_plus3.py
================
make_cv_plus2.py をベースに SNS・評判情報セクションを強化した実験版。
変更点:
  - SNS_INFO_PROMPT: 「確認できた情報のみ」制約を緩和し、
    学習データ・推論ベースの論評・推定評判も生成するよう指示を変更。
  - fetch_sns_info: max_tokens を 2500 に引き上げ。
  - それ以外は make_cv_plus2.py と完全に同一。

出力先: C:\\Users\\shondo\\Desktop\\agent_project\\cv
依存:   pip install anthropic python-docx pillow
"""

import io
import json
import os
import re
import sys
import urllib.request
import urllib.parse
from io import BytesIO
from pathlib import Path
from datetime import datetime

if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
    sys.stdin  = io.TextIOWrapper(sys.stdin.buffer,  encoding="utf-8", errors="replace")

try:
    import anthropic
except ImportError:
    sys.exit("anthropic が未インストールです。pip install anthropic を実行してください。")

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("python-docx が未インストールです。pip install python-docx を実行してください。")

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import fitz
except ImportError:
    fitz = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

# ── 設定 ─────────────────────────────────────────────────────────────────────
CLAUDE_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
CLAUDE_MODEL   = "claude-sonnet-4-6"
CV_DIR         = Path(r"C:\Users\shondo\Desktop\agent_project\cv")
DB_BASE        = Path(r"C:\Users\shondo\Desktop\agent_project\db")
F              = 12
HEADERS        = {"User-Agent": "Mozilla/5.0 (compatible; CVMaker/1.0)"}
MOFA_HEADERS   = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept": (
        "text/html,application/xhtml+xml,application/xml;q=0.9,"
        "image/avif,image/webp,*/*;q=0.8"
    ),
    "Accept-Language": "ja,en-US;q=0.9,en;q=0.8",
    "Accept-Encoding": "gzip, deflate, br",
    "Referer": "https://www.mofa.go.jp/",
    "Connection": "keep-alive",
}
GENERIC_HEADERS = {          # 大使館など外部サイト用（Referer なし）
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "ja,en-US;q=0.9,en;q=0.8",
    "Accept-Encoding": "gzip, deflate, br",
    "Connection": "keep-alive",
}
CHARS_PER_FILE = 4000

DB_FOLDERS = {
    "asean_oceania": "ASEAN諸国・オセアニア（タイ、ベトナム、インドネシア、シンガポール、豪州など）",
    "china":         "中国（習近平政権、中国共産党、人民元、香港など）",
    "europe":        "欧州（EU、英独仏伊、ウクライナ、ロシア、NATO欧州側など）",
    "geopolitics":   "地政学・地経学全般（半導体、サプライチェーン、鉱物資源など）",
    "india_mena":    "インド・中東・アフリカ（サウジ、イスラエル、イラン、UAE、ケニアなど）",
    "japan":         "日本（日本政府、日本企業、日本の外交・経済政策など）",
    "minerals":      "鉱物資源（レアアース、ニッケル、リチウム、黒鉛など）",
    "north america": "北米（米国、トランプ政権、カナダ、メキシコ、FRBなど）",
    "south america": "南米（ブラジル、アルゼンチン、ベネズエラなど）",
    "taiwan":        "台湾（台湾政府、台湾海峡、TSMC、台湾の安全保障など）",
    "trade_tariff_ai": "通商・関税・AI（WTO、FTA、人権、環境、AI規制など）",
    "others":        "上記以外",
}

# ── Claude プロンプト（名前解決・言語特定） ──────────────────────────────────
NAME_RESOLVE_PROMPT = """\
以下の人物について、日本語名・現地語名・情報検索に最適な言語を特定してください。

【入力情報】
入力された名前: {name}
国籍ヒント: {nationality_hint}
役職ヒント: {position_hint}

【出力ルール】
- JSON のみを出力（前置き・コードブロック不要）
- 確認できない情報は "不明" とする

{{
  "full_name_ja": "日本語・カタカナ表記の氏名（例: フランク・ルロワ）",
  "full_name_native": "現地語での正式氏名（例: Franck Leroy）。日本人なら日本語表記",
  "search_language": "情報検索に最適な言語の日本語名（例: フランス語、英語、ドイツ語、日本語）",
  "search_language_code": "BCP47言語コード（例: fr, en, de, ja）",
  "nationality": "国籍（例: フランス共和国）",
  "current_position": "現職（30文字以内）"
}}
"""

# ── Claude プロンプト（基本情報） ─────────────────────────────────────────────
INFO_PROMPT = """\
あなたは政治家・外交官・有識者の情報を整理する専門家です。
以下の人物について、日本の外交官が面談準備に使う略歴書用の情報を
JSON 形式で提供してください。

【対象人物】
氏名（日本語）: {name_ja}
氏名（現地語）: {name_native}
国籍: {nationality}
現職（ヒント）: {current_position}

【重要な検索指示】
現地語名「{name_native}」を使い、{search_language}でも情報を参照・検索してください。
日本語名だけでなく、現地語名での検索・参照を優先してください。

【出力ルール】
- JSON のみを出力（前置き・説明・コードブロック記号は一切不要）
- 各フィールドの文字数は必ず守ること（Word 文書の 1 行に収める制約）
- 情報が不確かな場合は "不明" と記載

{{
  "full_name_ja": "日本語・カタカナ表記の氏名",
  "birth_date": "生年月日（例: 1957年 4月 22日）",
  "nationality": "国籍（例: ポーランド共和国）",
  "current_position": "現職（30文字以内）",
  "education": [
    {{"period": "在籍期間（例: 1975年 ～ 1980年）",
      "description": "学校名・学部（25文字以内）"}}
  ],
  "career": [
    {{"year": "年（例: 2007年）",
      "description": "役職・出来事（25文字以内、現職・功績に関係するものに絞る、最大 8 件）"}}
  ],
  "interests": ["関心事項・主要政策（各28文字以内、5〜7件）"],
  "todo": ["面談で歓迎される話題・姿勢（各28文字以内、5〜7件）"],
  "not_todo": ["面談で避けるべき言動（各28文字以内、5〜7件）"],
  "wikipedia_title_en": "英語版 Wikipedia の記事タイトル（例: Donald Tusk）"
}}
"""

FOLDER_DETECT_PROMPT = """\
以下の人物が「最も関係する地域・テーマ」に対応するフォルダーを1つだけ選んでください。

【人物情報】
氏名: {name}
国籍: {nationality}
現職: {current_position}

【選択肢】
{folder_list}

回答はフォルダー名のみ1行で返してください（説明不要）。
"""

OFFICIAL_PHOTO_PROMPT = """\
以下の人物の顔写真（ポートレート・証明写真・バストショット）が掲載されている
公式ウェブサイトの画像 URL を最大 3 件提示してください。

【人物情報】
氏名（日本語）: {name_ja}
氏名（現地語）: {name_native}
現職: {current_position}
国籍: {nationality}

優先順位:
  1. 大使・外交官の場合 → 派遣先大使館の公式サイト（例: jp.usembassy.gov 等）
  2. 政府高官の場合    → 所属省庁・政府機関の公式サイト
  3. 国際機関職員     → 当該機関の公式サイト

厳守事項:
  - 必ず「人物の顔・上半身が写っている写真」の URL のみ
  - 建物・国旗・ロゴ・地図・バナー画像の URL は絶対に含めない
  - 直接アクセスできる画像ファイル（.jpg / .png / .webp）の URL のみ
  - 不明な場合は「不明」とだけ回答

URL のみを 1 行ずつ返してください（最大 3 件、説明不要）。
"""

AMBASSADOR_PROFILE_URL_PROMPT = """\
以下の大使について、公式プロフィールページのURLを最大5件提供してください。

【大使情報】
氏名（日本語）: {name}
種別: {ambassador_type}
相手国: {country}

【指示】
この大使の名前でウェブ検索すると見つかる公式プロフィールページ
（大使館公式サイトの大使紹介・プロフィールページ）のURLを
確率の高い順に列挙してください。
日本語版・英語版どちらも可。

URLのみを1行ずつ返してください（最大5件、説明不要）。
"""

EMBASSY_BASE_URL_PROMPT = """\
以下の大使館の公式ウェブサイトのトップページURLを1件返してください。

種別: {ambassador_type}
相手国: {country}
大使名（日本語）: {name}

URLのみ1行で返してください（説明不要）。
"""

NEWS_PARSE_PROMPT = """\
以下のウェブサイトから取得したテキストを分析して、
{name}（{country}大使）の最近の主な活動・イベントを
新しい順（最新が先頭）で3〜5件まとめてください。

【サイトテキスト】
{text}

【出力ルール】
- JSON のみを出力（前置き・コードブロック不要）
- 活動の日付・内容が明確なものだけを抽出
- 不明・確認できない場合は空リスト []

{{
  "recent_activities": [
    {{
      "date": "活動日・期間（例: 2025年5月）",
      "title": "活動名（25文字以内）",
      "purpose": "目的・背景（50文字以内）",
      "result": "成果・内容（50文字以内）"
    }}
  ],
  "activity_sources": ["情報源URL"]
}}
"""

ENRICH_PROMPT = """\
あなたは外交官向け面談準備資料の専門家です。
以下の人物の略歴書草案に、データベースファイルの最新情報を反映して
interests・todo・not_todo を更新してください。

【対象人物】
氏名（日本語）: {name_ja} / 現地語: {name_native}
国籍: {nationality}  現職: {current_position}

【データベースファイルの内容】
{files_block}

【現在の草案（interests / todo / not_todo）】
{current_sections}

【指示】
- データベースの情報を踏まえ、より具体的・最新の内容に更新する
- 草案の内容も活かしつつ、DB 情報で補強・修正する
- 各項目は28文字以内
- interests: 5〜7件、todo: 5〜7件、not_todo: 5〜7件
- JSON のみを出力（前置き不要）。形式:
{{"interests": [...], "todo": [...], "not_todo": [...]}}
"""

# ── Claude プロンプト（パーソナル情報） ──────────────────────────────────────
PERSONAL_INFO_PROMPT = """\
あなたは政治家・外交官・有識者に関する公開情報を調査する専門家です。
以下の人物について、公式ウェブサイト・公的プロフィール・インタビュー記事など
「公的・公開情報」をもとに、パーソナルな情報と懇談に役立つ話題を提供してください。

【対象人物】
氏名（日本語）: {name_ja}
氏名（現地語）: {name_native}
国籍: {nationality}
現職: {current_position}

【重要な検索指示】
現地語名「{name_native}」を使い、{search_language}でも情報を参照・検索してください。
現地語での公式サイト・インタビュー記事を優先して参照してください。

【出力ルール】
- JSON のみを出力（前置き・説明・コードブロック記号は一切不要）
- 確認できた情報のみ記載。不明な場合は "不明" または空リスト []
- 推測・伝聞は「〜とされる」「〜と伝えられる」と明記
- 各文字列は40文字以内

{{
  "family": [
    "家族構成・配偶者・子供など（公開情報のみ、各40文字以内）"
  ],
  "hobbies": [
    "趣味・関心（例: ゴルフ、読書、クラシック音楽）（各30文字以内）"
  ],
  "leisure": [
    "余暇の過ごし方（例: テニス、ハイキング、料理）（各30文字以内）"
  ],
  "pets": [
    "ペット情報（例: 犬2匹（ゴールデンレトリバー）など、不明な場合は 不明）"
  ],
  "personal_notes": [
    "その他パーソナルな公開情報（出身地、食の好み、信条など）（各40文字以内）"
  ],
  "conversation_topics": [
    "有識者懇談で盛り上がる可能性のある話題（各40文字以内、5〜7件）"
  ],
  "official_sources": [
    "情報源となった公式サイト・インタビューの媒体名またはURL"
  ]
}}
"""

# ── Claude プロンプト（SNS・評判情報）【強化版】────────────────────────────
SNS_INFO_PROMPT = """\
あなたはSNS・メディア上の評判調査の専門家です。
以下の人物について、X（旧Twitter）・Instagram・LinkedIn・Facebook・YouTube等の
各プラットフォームで「第三者が語っている評判・論評・エピソード・話題」を詳細に分析してください。

【対象人物】
氏名（日本語）: {name_ja}
氏名（現地語）: {name_native}
国籍: {nationality}
現職: {current_position}

【重要な検索指示】
現地語名「{name_native}」を使い、{search_language}のSNS・メディアも必ず参照してください。
フランス語・ドイツ語等の現地語での論評・報道も含めて分析してください。

【出力方針（重要）】
① 公式確認済みの情報があればそれを優先して記載する。
② 公式確認できない場合も、この人物の役職・政策・発言歴・公的立場をもとに
   「SNS上で語られているとみられる評判・論調」を必ず推定・記載すること。
   その場合は「〜と伝えられる」「〜との評価が多い」「〜とみられる」の形で明記する。
③ 各プラットフォーム（x_twitter・instagram・other_sns）は必ず最低2件以上記載する。
   「不明」「情報なし」のみの空配列は禁止。
④ media_reputation と notable_quotes は特に重要。必ず3〜5件の詳細な内容を記載する。
⑤ 誹謗中傷・根拠のない悪質な情報は除外し、公益性のある内容に限る。

【各プラットフォームの記載指針】
- x_twitter: 政治家・有識者・市民がこの人物の政策・言動・人柄についてXで語る論調を記載。
  国内外の評価・批判・支持の両面を含む。
- instagram: 本人または関係者のInstagram投稿に対する反応、地域密着活動・イベント写真等への論評。
  アカウントが確認できない場合は「公式アカウント不明だが、地域活動の投稿が話題とみられる」等の形で記載。
- other_sns: LinkedIn（職歴・経歴への評価）・Facebook（地元支持者の反応）等。
- media_reputation: 主要メディア・ジャーナリスト・政治評論家が語るこの人物の人物像・政治スタイル・評判。
- notable_quotes: 本人の発言・演説からの引用、または著名人・メディアがこの人物について述べた注目発言。

【出力ルール】
- JSON のみを出力（前置き・説明・コードブロック記号は一切不要）
- 各文字列は50文字以内（notable_quotes のみ60文字以内）

{{
  "x_twitter": [
    "Xでの評判・論評・話題（各50文字以内、3〜5件）"
  ],
  "instagram": [
    "Instagramでの評判・投稿内容・反応（各50文字以内、2〜3件）"
  ],
  "other_sns": [
    "LinkedIn・Facebook等での評判・論評（各50文字以内、2〜3件）"
  ],
  "media_reputation": [
    "メディア・評論家が語る人物像・評判（各50文字以内、3〜5件）"
  ],
  "notable_quotes": [
    "本人または第三者の注目発言・引用（各60文字以内、2〜3件）"
  ]
}}
"""


# ── Claude プロンプト（最近の活動） ──────────────────────────────────────────
RECENT_ACTIVITIES_PROMPT = """\
あなたは外交官・政治家・有識者の公開情報を調査する専門家です。
以下の人物について、大使館公式サイト・外務省・ニュース記事・Google検索結果をもとに
「最近の主な活動」を時系列で調査してください。

【対象人物】
氏名（日本語）: {name_ja}
氏名（現地語）: {name_native}
国籍: {nationality}
現職: {current_position}

【重要な検索指示】
現地語名「{name_native}」を使い、{search_language}でも情報を参照・検索してください。
現地語での公式サイト・プレスリリース・ニュース記事を優先してください。

【調査対象】
- 大使館公式ウェブサイト・プレスリリース・ニュースリリース
- 外務省・相手国政府の発表・公式記録
- 日本国内の報道（NHK、共同通信、各新聞社等）
- Google 等での検索結果（公式情報・信頼性の高いメディアに限る）

【出力ルール】
- JSON のみを出力（前置き・コードブロック不要）
- 直近1〜2年の活動を優先し、新しい順（最新が先頭）で3〜5件
- 確認できた事実のみ記載。推測には「〜とみられる」と明記
- 不明・確認できない場合は空リスト []
- title は25文字以内、purpose・result は各50文字以内

{{
  "recent_activities": [
    {{
      "date": "活動日・期間（例: 2024年11月）",
      "title": "活動名・イベント名（25文字以内、例: 北海道農業・食品産業視察）",
      "purpose": "目的・背景（50文字以内、例: 農業分野の二国間協力強化と輸出促進）",
      "result": "成果・結果（50文字以内、例: 道内企業と意見交換、農産品輸出拡大で合意）"
    }}
  ],
  "activity_sources": [
    "情報源となったサイトURL または媒体名（確認できた範囲で）"
  ]
}}
"""


# ── Claude プロンプト（大使クエリ判定） ──────────────────────────────────────
AMBASSADOR_DETECT_PROMPT = """\
以下の入力文字列が「大使または外交官」に関する調査依頼かどうかを判断してください。

【入力】{query}

【判断基準】
- 「チェコ大使」「駐日フランス大使」「特命全権大使」などが含まれる場合は大使関連
- 人名が入力された場合、その人物が大使・外交官として著名であれば大使関連
- それ以外（政治家・経済人・研究者など）は大使関連でない

【ambassador_type の定義】
  - "駐日大使": 相手国が日本に派遣している大使（例: 駐日チェコ大使）
  - "駐外大使": 日本が相手国に派遣している大使（例: 駐チェコ日本大使）
  - "不明": 種別が判断できない場合

【出力ルール】
- JSON のみを出力（前置き・コードブロック不要）

{{
  "is_ambassador": true または false,
  "ambassador_type": "駐日大使" または "駐外大使" または "不明",
  "country_name": "国名（例: チェコ、フランス、アメリカ）。不明な場合は空文字",
  "mofa_page_url": "外務省の当該国ページURL（例: https://www.mofa.go.jp/mofaj/area/czech/index.html）。不明な場合は空文字"
}}
"""


# ── XML / フォントヘルパー ─────────────────────────────────────────────────────
def _set_east_asian_font(run, font_name: str):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "FFFFFF")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def set_cell_width(cell, width_cm: float):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")


def set_cell_font(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(F)
    run.font.bold = bold
    run.font.name = "MS明朝"
    _set_east_asian_font(run, "MS明朝")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


# ── 文書生成ヘルパー ──────────────────────────────────────────────────────────
def para(doc, text: str = "", size: int = F, bold: bool = False,
         align=WD_ALIGN_PARAGRAPH.LEFT, color=None,
         space_after: int = 3, underline: bool = False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.size      = Pt(size)
        run.font.bold      = bold
        run.font.underline = underline
        run.font.name      = "MS明朝"
        _set_east_asian_font(run, "MS明朝")
        if color:
            run.font.color.rgb = color
    return p


def add_info_line(cell, label: str, value: str, first: bool = False):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(label)
    r1.font.size = Pt(F)
    r1.font.name = "MS明朝"
    _set_east_asian_font(r1, "MS明朝")
    r2 = p.add_run(f"　{value}")
    r2.font.size = Pt(F)
    r2.font.bold = True
    r2.font.name = "MS明朝"
    _set_east_asian_font(r2, "MS明朝")


def add_bullets(doc, items: list):
    for item in items:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after  = Pt(3)
        run = bp.add_run(item)
        run.font.size = Pt(F)
        run.font.name = "MS明朝"
        _set_east_asian_font(run, "MS明朝")


def add_section_header(doc, text: str):
    """区切り線付きセクションヘッダー"""
    para(doc, text, size=F, bold=True, space_after=4, underline=True)


def add_kv_table(doc, rows_data: list):
    """
    2列の Key-Value テーブルを追加する。
    rows_data = [("ラベル", "値"), ...]
    """
    if not rows_data:
        return
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style   = "Table Grid"
    tbl.autofit = False
    for i, (label, value) in enumerate(rows_data):
        set_cell_font(tbl.rows[i].cells[0], label, bold=True)
        set_cell_font(tbl.rows[i].cells[1], value)
        set_cell_width(tbl.rows[i].cells[0], 3.8)
        set_cell_width(tbl.rows[i].cells[1], 12.7)
    remove_table_borders(tbl)


def add_sns_block(doc, label: str, items: list):
    """SNSプラットフォーム名ラベル＋箇条書きブロック"""
    if not items or items == ["不明"]:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(f"［{label}］")
    run.font.size = Pt(F)
    run.font.bold = True
    run.font.name = "MS明朝"
    run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
    _set_east_asian_font(run, "MS明朝")
    add_bullets(doc, items)


# ── 名前解決・言語特定 ────────────────────────────────────────────────────────
def resolve_name_multilingual(name: str, nationality_hint: str,
                               position_hint: str, client) -> dict:
    """入力名・国籍ヒント・役職ヒントから日本語名・現地語名・検索言語を特定する。"""
    prompt = NAME_RESOLVE_PROMPT.format(
        name=name,
        nationality_hint=nationality_hint or "不明",
        position_hint=position_hint or "不明",
    )
    try:
        msg = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=300,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip().lstrip("﻿")
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
        raw = re.sub(r"\s*```\s*$", "", raw)
        return json.loads(raw)
    except Exception as e:
        print(f"  [名前解決] エラー: {e}  入力をそのまま使用します。")
        return {
            "full_name_ja": name,
            "full_name_native": name,
            "search_language": "日本語",
            "search_language_code": "ja",
            "nationality": nationality_hint or "不明",
            "current_position": position_hint or "不明",
        }


# ── Claude で人物情報を取得 ────────────────────────────────────────────────────
def fetch_person_info(name_ja: str, client, name_native: str = "",
                      nationality: str = "", current_position: str = "",
                      search_language: str = "日本語") -> dict:
    content = INFO_PROMPT.format(
        name_ja=name_ja,
        name_native=name_native or name_ja,
        nationality=nationality or "不明",
        current_position=current_position or "不明",
        search_language=search_language,
    )
    msg = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=2000,
        messages=[{"role": "user", "content": content}],
    )
    raw = msg.content[0].text.strip().lstrip("﻿")
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s*```\s*$", "", raw)
    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"  JSON パースエラー: {e}")
        print(f"  出力の先頭 300 文字:\n{raw[:300]}")
        raise


# ── DB フォルダーの判定 ────────────────────────────────────────────────────────
def detect_db_folder(info: dict, client) -> Path | None:
    available = {k: v for k, v in DB_FOLDERS.items()
                 if (DB_BASE / k).is_dir()}
    if not available:
        return None
    folder_list = "\n".join(f"  {k}: {v}" for k, v in available.items())
    prompt = FOLDER_DETECT_PROMPT.format(
        name=info.get("full_name_ja", ""),
        nationality=info.get("nationality", ""),
        current_position=info.get("current_position", ""),
        folder_list=folder_list,
    )
    msg = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=30,
        messages=[{"role": "user", "content": prompt}],
    )
    chosen = msg.content[0].text.strip().lower().strip('"').strip("'")
    for key in available:
        if key in chosen:
            folder = DB_BASE / key
            print(f"  [DB] 関連フォルダー: {folder}")
            return folder
    return None


# ── DB ファイルのテキスト抽出 ─────────────────────────────────────────────────
def _extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            if pdfplumber:
                with pdfplumber.open(str(path)) as pdf:
                    return "\n".join(p.extract_text() or "" for p in pdf.pages)
            if fitz:
                doc = fitz.open(str(path))
                return "\n".join(page.get_text() for page in doc)
            return ""
        if ext == ".docx":
            doc = Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        if ext in (".txt", ".md"):
            return path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        pass
    return ""


def build_files_block(folder: Path) -> str:
    supported = {".pdf", ".docx", ".txt", ".md"}
    files = sorted(folder.rglob("*"))
    blocks = []
    for f in files:
        if f.suffix.lower() not in supported:
            continue
        text = _extract_text(f)
        if not text.strip():
            continue
        snippet = text[:CHARS_PER_FILE]
        if len(text) > CHARS_PER_FILE:
            snippet += "\n[以下省略]"
        blocks.append(f"=== {f.name} ===\n{snippet}")
        if len(blocks) >= 15:
            break
    return "\n\n".join(blocks)


# ── DB 情報で interests / todo / not_todo を補強 ──────────────────────────────
def enrich_with_db(info: dict, files_block: str, client,
                   name_native: str = "") -> dict:
    current_sections = json.dumps({
        "interests": info.get("interests", []),
        "todo":      info.get("todo", []),
        "not_todo":  info.get("not_todo", []),
    }, ensure_ascii=False, indent=2)
    prompt = ENRICH_PROMPT.format(
        name_ja=info.get("full_name_ja", ""),
        name_native=name_native or info.get("full_name_ja", ""),
        nationality=info.get("nationality", ""),
        current_position=info.get("current_position", ""),
        files_block=files_block,
        current_sections=current_sections,
    )
    msg = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = msg.content[0].text.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s*```\s*$", "", raw)
    try:
        enriched = json.loads(raw)
        for key in ("interests", "todo", "not_todo"):
            if key in enriched and isinstance(enriched[key], list):
                info[key] = enriched[key]
        return info
    except json.JSONDecodeError as e:
        print(f"  [DB補強] JSON パースエラー: {e}  草案のまま使用します。")
        return info


# ── 大使館写真を取得（名前検索→候補URL順試し） ─────────────────────────────
def _try_photo_from_page(page_url: str, name_parts: list, save_path: Path) -> bool:
    """
    指定URLのページをスクレイピングし、大使のポートレート写真をダウンロードする。

    改善点:
    - "velvyslanec" を単語境界 \\b で検索し、"velvyslanectví"（大使館）を誤マッチしない
    - src に建物系キーワードが含まれる画像を除外
    - 全候補をスコアリングし、ダウンロード後にアスペクト比チェック（横長=建物バナーを除外）
    - 小さすぎる画像（アイコン等）を除外
    成功したら True を返す。
    """
    try:
        html = _fetch_html_generic(page_url)
    except Exception as e:
        print(f"  [大使写真] ページ取得失敗: {e}")
        return False

    # src パスに含まれているとスキップするキーワード（建物・ロゴ等）
    SKIP_SRC_KW = (
        "building", "exterior", "facade", "budova", "budynek",
        "logo", "icon", "banner", "flag", "map", "background",
        "coat", "crest", "arrow", "button",
    )

    candidates = []  # (score, photo_url, alt_text)

    for m in re.finditer(r'<img([^>]+)>', html, re.IGNORECASE):
        tag = m.group(1)
        src_m = re.search(r'src=["\']([^"\']+)["\']', tag, re.I)
        alt_m = re.search(r'alt=["\']([^"\']*)["\']', tag, re.I)
        if not src_m:
            continue
        src = src_m.group(1)
        alt = alt_m.group(1) if alt_m else ""

        # 画像ファイル拡張子チェック
        if not any(src.lower().split("?")[0].endswith(e)
                   for e in (".jpg", ".jpeg", ".png", ".webp")):
            continue

        # ネガティブフィルター: src に建物・ロゴ系キーワードを含む画像をスキップ
        if any(kw in src.lower() for kw in SKIP_SRC_KW):
            continue

        # スコアリング
        alt_low = alt.lower()
        score = 0
        # "ambassador" / "velvyslanec" を単語境界で検索
        # （"velvyslanectví" = 大使館 を誤マッチしないため \b を使用）
        if re.search(r'\bambassador\b', alt_low):
            score += 10
        if re.search(r'\bvelvyslanec\b', alt_low):
            score += 10
        # 大使名の一部が含まれる
        for p in name_parts:
            if len(p) > 2 and p.lower() in alt_low:
                score += 5

        if score > 0:
            photo_url = _abs_url(src, page_url)
            candidates.append((score, photo_url, alt))

    if not candidates:
        return False

    # スコアの高い順に試す
    candidates.sort(key=lambda x: -x[0])

    for score, photo_url, alt in candidates:
        print(f"  [大使写真] 候補(score={score}): {photo_url}  (alt={alt[:80]})")
        try:
            req = urllib.request.Request(photo_url, headers=GENERIC_HEADERS)
            with urllib.request.urlopen(req, timeout=10) as resp:
                img_data = resp.read()

            # アスペクト比・サイズチェック（PIL がある場合）
            if Image:
                img = Image.open(BytesIO(img_data))
                w, h = img.size
                # 横長すぎる画像はスキップ（建物・バナー等; 幅 > 高さ × 1.5）
                if w > h * 1.5:
                    print(f"  [大使写真] スキップ（横長画像 {w}x{h}）: {photo_url}")
                    continue
                # 小さすぎる画像はスキップ（アイコン等）
                if w < 50 or h < 50:
                    print(f"  [大使写真] スキップ（小さすぎる {w}x{h}）: {photo_url}")
                    continue
                img.convert("RGB").save(str(save_path), "JPEG", quality=90)
            else:
                save_path.write_bytes(img_data)

            print(f"  [大使写真] 保存しました: {save_path.name}")
            return True
        except Exception as e:
            print(f"  [大使写真] ダウンロードエラー: {e}")

    return False


def _try_photo_with_sublinks(page_url: str, name_parts: list, save_path: Path) -> bool:
    """
    指定URLで写真を探す。
    見つからず、かつ index.html 系のページの場合は同ディレクトリの
    サブリンクを1段階追跡してさらに探す。
    """
    if _try_photo_from_page(page_url, name_parts, save_path):
        return True

    # index / ディレクトリページでなければここで終了
    clean = page_url.split("?")[0].lower()
    if not (clean.endswith("index.html") or clean.endswith("index.htm")
            or clean.endswith("/")):
        return False

    print(f"  [大使写真] サブリンクを探索: {page_url}")
    try:
        html = _fetch_html_generic(page_url)
    except Exception:
        return False

    dir_prefix = page_url.rsplit("/", 1)[0] + "/"
    seen = {page_url}
    for href in _extract_hrefs(html):
        abs_href = _abs_url(href, page_url)
        if abs_href in seen:
            continue
        if not abs_href.startswith(dir_prefix):
            continue
        seen.add(abs_href)
        last = abs_href.split("/")[-1]
        if "#" in last or "?" in last:
            continue
        if "." in last:
            ext = last.rsplit(".", 1)[-1].lower()
            if ext in ("svg", "png", "jpg", "gif", "css", "js",
                       "pdf", "ico", "xml", "zip"):
                continue
        print(f"  [大使写真] サブリンク確認: {abs_href}")
        if _try_photo_from_page(abs_href, name_parts, save_path):
            return True

    return False


def fetch_ambassador_photo(info: dict, ambassador_type: str, country: str,
                           client, save_path: Path,
                           embassy_pages: dict = None) -> bool:
    """
    ① Claude に大使名でたどり着けるプロフィールページURLを最大5件問い合わせ
    ② 各URLをスクレイピングして写真を探す
    ③ 全部失敗なら大使館サイトナビゲーション（embassy_pages）でフォールバック
    """
    full_name = info.get("full_name_ja", "")
    name_parts = [p for p in full_name.split() if len(p) > 1]

    # ── Step1: Claude にプロフィールページ候補URLを問い合わせ ──────────────────
    prompt = AMBASSADOR_PROFILE_URL_PROMPT.format(
        name=full_name,
        ambassador_type=ambassador_type or "大使",
        country=country,
    )
    try:
        msg = client.messages.create(
            model=CLAUDE_MODEL, max_tokens=300,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        candidate_urls = [
            line.strip() for line in raw.splitlines()
            if line.strip().startswith("http")
        ]
        print(f"  [大使写真] Claude が返した候補URL: {len(candidate_urls)} 件")
        for u in candidate_urls:
            print(f"    {u}")
    except Exception as e:
        print(f"  [大使写真] 候補URL取得エラー: {e}")
        candidate_urls = []

    # ── Step2: 候補URLを順番に試す（index ページはサブリンクも1段追跡） ──────────
    for page_url in candidate_urls:
        print(f"  [大使写真] ページを確認: {page_url}")
        if _try_photo_with_sublinks(page_url, name_parts, save_path):
            return True

    # ── Step3: embassy_pages のナビゲーション結果でフォールバック ───────────────
    if embassy_pages:
        for key in ("ambassador_page", "embassy_base"):
            page_url = embassy_pages.get(key, "")
            if not page_url:
                continue
            print(f"  [大使写真] フォールバック確認: {page_url}")
            if _try_photo_with_sublinks(page_url, name_parts, save_path):
                return True

    print("  [大使写真] 写真を取得できませんでした。")
    return False


# ── Claude ビジョンで顔写真かどうかを検証 ─────────────────────────────────────
def _is_portrait_photo(img_data: bytes, client) -> bool:
    """ダウンロードした画像をClaudeに見せて人物の顔・上半身が写っているか確認する。"""
    import base64
    try:
        if Image:
            img = Image.open(BytesIO(img_data))
            fmt = (img.format or "JPEG").upper()
            media_type = {
                "JPEG": "image/jpeg", "JPG": "image/jpeg",
                "PNG":  "image/png",  "WEBP": "image/webp",
                "GIF":  "image/gif",
            }.get(fmt, "image/jpeg")
        else:
            media_type = "image/jpeg"

        img_b64 = base64.standard_b64encode(img_data).decode("utf-8")
        msg = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=10,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image",
                     "source": {"type": "base64",
                                "media_type": media_type,
                                "data": img_b64}},
                    {"type": "text",
                     "text": "この画像に人物の顔または上半身が写っていますか？「はい」か「いいえ」のみで答えてください。"},
                ]
            }],
        )
        answer = msg.content[0].text.strip()
        return "はい" in answer
    except Exception as e:
        print(f"  [写真検証] エラー: {e}  → 通過扱い")
        return True


# ── 公式サイトから写真を取得 ──────────────────────────────────────────────────
def fetch_official_photo(info: dict, client, save_path: Path,
                         name_native: str = "") -> bool:
    prompt = OFFICIAL_PHOTO_PROMPT.format(
        name_ja=info.get("full_name_ja", ""),
        name_native=name_native or info.get("full_name_ja", ""),
        current_position=info.get("current_position", ""),
        nationality=info.get("nationality", ""),
    )
    try:
        msg = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=300,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
    except Exception as e:
        print(f"  [公式写真] URL取得エラー: {e}")
        return False

    candidate_urls = [
        line.strip() for line in raw.splitlines()
        if line.strip().startswith("http")
    ]
    if not candidate_urls:
        print("  [公式写真] 公式サイトの写真URLが見つかりませんでした。")
        return False

    for url in candidate_urls:
        print(f"  [公式写真] 取得中: {url}")
        try:
            req = urllib.request.Request(url, headers=HEADERS)
            with urllib.request.urlopen(req, timeout=10) as resp:
                img_data = resp.read()
        except Exception as e:
            print(f"  [公式写真] 取得エラー: {e}")
            continue

        # Claude ビジョンで顔・上半身が写っているか確認
        print(f"  [公式写真] 顔写真かどうかを検証中...")
        if not _is_portrait_photo(img_data, client):
            print(f"  [公式写真] スキップ（顔・上半身なし）: {url}")
            continue

        if Image:
            img = Image.open(BytesIO(img_data)).convert("RGB")
            img.save(str(save_path), "JPEG", quality=90)
        else:
            save_path.write_bytes(img_data)

        print(f"  [公式写真] 保存しました: {save_path.name}")
        return True

    print("  [公式写真] すべての候補URLで取得できませんでした。")
    return False


# ── Wikipedia から写真を取得 ──────────────────────────────────────────────────
def fetch_wikipedia_photo(wikipedia_title: str, save_path: Path) -> bool:
    if not wikipedia_title or wikipedia_title == "不明":
        return False
    try:
        encoded = urllib.parse.quote(wikipedia_title.replace(" ", "_"))
        api_url = (
            "https://en.wikipedia.org/w/api.php"
            f"?action=query&titles={encoded}&prop=pageimages"
            "&format=json&pithumbsize=400"
        )
        req = urllib.request.Request(api_url, headers=HEADERS)
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = json.loads(resp.read())
        pages = data.get("query", {}).get("pages", {})
        thumb_url = None
        for page in pages.values():
            thumb_url = (page.get("thumbnail") or {}).get("source")
            break
        if not thumb_url:
            print("  [写真] Wikipedia に写真が見つかりませんでした。写真なしで作成します。")
            return False
        print("  [写真] 取得中 ...")
        req2 = urllib.request.Request(thumb_url, headers=HEADERS)
        with urllib.request.urlopen(req2, timeout=10) as resp:
            img_data = resp.read()
        if Image:
            img = Image.open(BytesIO(img_data)).convert("RGB")
            img.save(str(save_path), "JPEG", quality=90)
        else:
            save_path.write_bytes(img_data)
        print(f"  [写真] 保存しました: {save_path.name}")
        return True
    except Exception as e:
        print(f"  [写真] 取得エラー: {e}  写真なしで作成します。")
        return False


# ── パーソナル情報を取得（公式ソース） ───────────────────────────────────────
def fetch_personal_info(info: dict, client, name_native: str = "",
                        search_language: str = "日本語") -> dict:
    """公式情報をもとに家族・趣味・余暇・ペット・懇談話題を取得する。"""
    prompt = PERSONAL_INFO_PROMPT.format(
        name_ja=info.get("full_name_ja", ""),
        name_native=name_native or info.get("full_name_ja", ""),
        nationality=info.get("nationality", ""),
        current_position=info.get("current_position", ""),
        search_language=search_language,
    )
    try:
        msg = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=1500,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
        raw = re.sub(r"\s*```\s*$", "", raw)
        return json.loads(raw)
    except Exception as e:
        print(f"  [パーソナル情報] 取得エラー: {e}  スキップします。")
        return {}


# ── 外務省スクレイピング（大使情報解決） ──────────────────────────────────────
def _fetch_mofa_html(url: str) -> str:
    """外務省ページを取得して UTF-8 文字列で返す（gzip 対応）"""
    import gzip as _gzip
    req = urllib.request.Request(url, headers=MOFA_HEADERS)
    with urllib.request.urlopen(req, timeout=15) as resp:
        raw = resp.read()
    try:
        return _gzip.decompress(raw).decode("utf-8", errors="replace")
    except Exception:
        return raw.decode("utf-8", errors="replace")


def _fetch_html_generic(url: str) -> str:
    """大使館など任意のサイトをgzip対応で取得する"""
    import gzip as _gzip
    req = urllib.request.Request(url, headers=GENERIC_HEADERS)
    with urllib.request.urlopen(req, timeout=15) as resp:
        raw  = resp.read()
        charset = resp.headers.get_content_charset() or "utf-8"
    try:
        decoded = _gzip.decompress(raw)
    except Exception:
        decoded = raw
    for enc in (charset, "utf-8", "latin-1"):
        try:
            return decoded.decode(enc, errors="replace")
        except Exception:
            pass
    return decoded.decode("utf-8", errors="replace")


def _abs_url(src: str, base_url: str) -> str:
    """相対URLを絶対URLに変換"""
    if src.startswith("http"):
        return src
    if src.startswith("//"):
        return "https:" + src
    from urllib.parse import urlparse
    p = urlparse(base_url)
    if src.startswith("/"):
        return f"{p.scheme}://{p.netloc}{src}"
    return base_url.rsplit("/", 1)[0] + "/" + src


def _extract_hrefs(html: str) -> list:
    """HTMLから全hrefを抽出してリストで返す"""
    return re.findall(r'href=["\']([^"\']+)["\']', html, re.IGNORECASE)


def _find_ambassador_page(hrefs: list, base_url: str) -> str:
    """href リストから大使紹介ページURLを探す（多言語対応）"""
    bad = ("svg","png","jpg","css","js","ico","pdf","#","?","mail","tel")
    # 英語・チェコ語・仏語・独語・西語・伊語・波語等の「大使」相当語を網羅
    AMB_WORDS = (
        "ambassador",    # 英語
        "velvyslanec",   # チェコ語
        "ambassadeur",   # 仏語
        "botschafter",   # 独語
        "embajador",     # 西語
        "ambasciatore",  # 伊語
        "ambasador",     # 波語・ルーマニア語等
        "ambassadör",    # スウェーデン語等
    )
    for lk in hrefs:
        lk_low = lk.lower()
        if any(w in lk_low for w in AMB_WORDS):
            if not any(lk_low.endswith(x) for x in bad):
                return _abs_url(lk, base_url)
    return ""


def _find_news_page(hrefs: list, base_url: str) -> str:
    """href リストからニュース/活動ページURLを探す（index.html優先）"""
    bad = ("svg","png","jpg","css","js","ico","pdf","#","?","mail","tel")
    keywords = ("bilateral","news","press","activity","event","relations","aktualne","aktual")
    candidates = []
    for lk in hrefs:
        if any(k in lk.lower() for k in keywords):
            if not any(lk.lower().endswith(x) for x in bad):
                candidates.append(lk)
    if not candidates:
        return ""
    candidates.sort(key=lambda x: (0 if "index" in x.lower() else 1, len(x)))
    return _abs_url(candidates[0], base_url)


def get_embassy_pages(ambassador_type: str, country: str, name: str, client) -> dict:
    """
    大使館トップURLをClaudeに問い合わせ、2段階でナビリンクを解析して
    大使紹介ページ・ニュースページのURLを返す。
    Returns: {"embassy_base": ..., "ambassador_page": ..., "news_page": ...}
    """
    # Step1: Claude にベースURL問い合わせ
    prompt = EMBASSY_BASE_URL_PROMPT.format(
        ambassador_type=ambassador_type or "大使",
        country=country,
        name=name,
    )
    try:
        msg = client.messages.create(
            model=CLAUDE_MODEL, max_tokens=100,
            messages=[{"role": "user", "content": prompt}],
        )
        base_url = msg.content[0].text.strip()
        if not base_url.startswith("http"):
            return {}
    except Exception as e:
        print(f"  [大使館] ベースURL取得エラー: {e}")
        return {}

    print(f"  [大使館] サイトを解析中: {base_url}")

    # Step2: トップページ取得
    try:
        html_top = _fetch_html_generic(base_url)
    except Exception as e:
        print(f"  [大使館] サイト取得失敗: {e}")
        return {"embassy_base": base_url}

    hrefs_top = _extract_hrefs(html_top)
    result = {"embassy_base": base_url}

    # ── 大使紹介ページを探索（2段階） ─────────────────────────────────────────
    # 第1段階: トップページ直接
    amb_page = _find_ambassador_page(hrefs_top, base_url)

    # 第2段階: サブセクションを1段掘り下げる（多言語対応）
    if not amb_page:
        sub_keywords = (
            "embassy", "about", "staff", "mission", "consulate", "diplomatic",
            "velvyslanectvi", "o_velvyslanectvi",   # チェコ語（大使館）
            "embajada", "botschaft", "ambasada", "ambassade",  # 他言語
        )
        sub_candidates = []
        for lk in hrefs_top:
            if any(k in lk.lower() for k in sub_keywords):
                if not any(lk.lower().endswith(x) for x in ("svg","png","jpg","css","js","#","?")):
                    sub_candidates.append(_abs_url(lk, base_url))
        # 重複除去・短いURL優先（index.html を含むもの優先）
        seen = set()
        sub_unique = []
        for u in sorted(sub_candidates, key=lambda x: (0 if "index" in x.lower() else 1, len(x))):
            if u not in seen:
                seen.add(u)
                sub_unique.append(u)

        for sub_url in sub_unique[:4]:   # 最大4ページまで掘り下げ
            print(f"  [大使館] サブページを確認: {sub_url}")
            try:
                html_sub = _fetch_html_generic(sub_url)
                hrefs_sub = _extract_hrefs(html_sub)
                amb_page = _find_ambassador_page(hrefs_sub, sub_url)
                if amb_page:
                    break
            except Exception:
                continue

    if amb_page:
        result["ambassador_page"] = amb_page
        print(f"  [大使館] 大使紹介ページ: {amb_page}")
    else:
        print("  [大使館] 大使紹介ページが見つかりませんでした。")

    # ── ニュースページ ──────────────────────────────────────────────────────────
    news_page = _find_news_page(hrefs_top, base_url)
    if news_page:
        result["news_page"] = news_page
        print(f"  [大使館] ニュースページ: {news_page}")

    return result


def scrape_mofa_ambassador_name(mofa_page_url: str, ambassador_type: str):
    """
    外務省の国データページ（data.html）から現在の大使名を取得する。

    ambassador_type:
      "駐日大使" → 相手国が日本に派遣している大使
      "駐外大使" → 日本が相手国に派遣している大使

    Returns: 大使名（文字列）または None
    """
    if not mofa_page_url:
        return None

    # index.html → data.html に変換
    data_url = re.sub(r'index\.html$', 'data.html', mofa_page_url)
    if data_url == mofa_page_url:
        data_url = mofa_page_url.rstrip("/") + "/data.html"

    print(f"  [MOFA大使] データページを取得中: {data_url}")
    try:
        html = _fetch_mofa_html(data_url)
    except Exception as e:
        print(f"  [MOFA大使] データページ取得失敗: {e}")
        return None

    # 「外交使節」見出し以降を対象にする
    idx = html.find("外交使節")
    if idx == -1:
        print("  [MOFA大使] 外交使節セクションが見つかりませんでした。")
        return None

    section = html[idx: idx + 800]

    # <li>…</li> を全件取得（class="none" 等の属性付きタグにも対応）
    items = re.findall(r'<li[^>]*>(.*?)</li>', section, re.DOTALL)
    for item in items:
        text = re.sub(r'<[^>]+>', '', item).strip()
        text = re.sub(r'^[（(]\d+[）)]\s*', '', text).strip()
        if not text:
            continue

        if ambassador_type == "駐日大使" and "駐日" in text:
            print(f"  [MOFA大使] 駐日大使を検出: {text}")
            return text
        elif ambassador_type == "駐外大使" and "駐日" not in text:
            print(f"  [MOFA大使] 駐外大使を検出: {text}")
            return text

    # 種別不明の場合は最初の項目を返す
    if ambassador_type == "不明" and items:
        text = re.sub(r'<[^>]+>', '', items[0]).strip()
        text = re.sub(r'^[（(]\d+[）)]\s*', '', text).strip()
        if text:
            print(f"  [MOFA大使] 大使名（種別不明）: {text}")
            return text

    print("  [MOFA大使] 大使情報を抽出できませんでした。")
    return None


def detect_ambassador_query(query: str, client):
    """
    入力クエリが大使関連かどうかを判定し、
    大使関連の場合は外務省 data.html から現在の大使名を取得して返す。

    Returns: (resolved_name, country_name, ambassador_type)
      resolved_name    : MOFAから取得した大使名（取得できなければ元のquery）
      country_name     : 国名（大使関連でなければ空文字）
      ambassador_type  : "駐日大使" / "駐外大使" / "不明" / ""
    """
    prompt = AMBASSADOR_DETECT_PROMPT.format(query=query)
    try:
        msg = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=400,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
        raw = re.sub(r"\s*```\s*$", "", raw)
        result = json.loads(raw)
    except Exception as e:
        print(f"  [大使判定] エラー: {e}  スキップします。")
        return query, "", ""

    if not result.get("is_ambassador", False):
        return query, "", ""

    country  = result.get("country_name", "")
    mofa_url = result.get("mofa_page_url", "")
    amb_type = result.get("ambassador_type", "不明")

    print(f"  [大使判定] 大使クエリを検出: 国={country}, 種別={amb_type}")
    print(f"  [大使判定] 外務省URL: {mofa_url}")

    if not mofa_url:
        print("  [大使判定] 外務省URLが不明のため元のクエリを使用します。")
        return query, country, amb_type

    ambassador_name = scrape_mofa_ambassador_name(mofa_url, amb_type)
    if ambassador_name:
        print(f"  [大使判定] 現在の大使名を確認: 「{ambassador_name}」")
        return ambassador_name, country, amb_type
    else:
        print("  [大使判定] 大使名を取得できませんでした。元のクエリを使用します。")
        return query, country, amb_type


# ── SNS・評判情報を取得【強化版】────────────────────────────────────────────
def fetch_sns_info(info: dict, client, name_native: str = "",
                   search_language: str = "日本語") -> dict:
    """X・Instagram等のSNS上で第三者が語っている評判・話題を取得する（強化版）。"""
    prompt = SNS_INFO_PROMPT.format(
        name_ja=info.get("full_name_ja", ""),
        name_native=name_native or info.get("full_name_ja", ""),
        nationality=info.get("nationality", ""),
        current_position=info.get("current_position", ""),
        search_language=search_language,
    )
    try:
        msg = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=2500,   # 1500 → 2500 に引き上げ（SNS各プラットフォーム分を確保）
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
        raw = re.sub(r"\s*```\s*$", "", raw)
        result = json.loads(raw)
        # プロンプト例文がそのまま出力された行を除去
        for key in ("x_twitter", "instagram", "other_sns", "media_reputation", "notable_quotes"):
            val = result.get(key, [])
            result[key] = [
                v for v in val
                if v and v != "不明"
                and "各50文字以内" not in v
                and "各60文字以内" not in v
            ]
        return result
    except Exception as e:
        print(f"  [SNS情報] 取得エラー: {e}  スキップします。")
        return {}


# ── 最近の活動を取得（スクレイピング優先） ───────────────────────────────────
def fetch_recent_activities(info: dict, client,
                            ambassador_country: str = "",
                            ambassador_type: str = "",
                            embassy_pages: dict = None,
                            name_native: str = "",
                            search_language: str = "日本語") -> dict:
    """
    大使クエリの場合は大使館ニュースページを実際にスクレイピングして構造化。
    それ以外は Claude の学習データから生成。
    """
    scraped_text = ""
    news_url     = ""

    # ── 大使館サイトからニュースを取得 ────────────────────────────────────────
    if ambassador_country and embassy_pages:
        news_url = embassy_pages.get("news_page", "")

        # embassy_pages にニュースURLがない場合、ベースから再探索
        if not news_url:
            base_url = embassy_pages.get("embassy_base", "")
            if base_url:
                try:
                    html_base = _fetch_html_generic(base_url)
                    for m in re.finditer(r'href=["\']([^"\'#?][^"\']*)["\']', html_base):
                        lk = m.group(1)
                        if any(k in lk.lower()
                               for k in ["bilateral","news","press","activity","event"]):
                            if not lk.lower().endswith(("svg","png","jpg","css","js")):
                                news_url = _abs_url(lk, base_url)
                                break
                except Exception:
                    pass

        if news_url and news_url.startswith("http"):
            print(f"  [最近の活動] ニュースページを取得中: {news_url}")
            try:
                html = _fetch_html_generic(news_url)
                # HTMLタグを除去してテキスト化
                text = re.sub(r'<[^>]+>', ' ', html)
                text = re.sub(r'&[a-z]+;', ' ', text)   # HTMLエンティティ除去
                text = re.sub(r'\s+', ' ', text).strip()
                scraped_text = text[:5000]
                print(f"  [最近の活動] {len(scraped_text)} 文字スクレイピング完了")
            except Exception as e:
                print(f"  [最近の活動] ニュースページ取得失敗: {e}")

    # ── プロンプト選択 ─────────────────────────────────────────────────────────
    if scraped_text:
        prompt = NEWS_PARSE_PROMPT.format(
            name=info.get("full_name_ja", ""),
            country=ambassador_country or info.get("nationality", ""),
            text=scraped_text,
        )
    else:
        # 大使館スクレイピング失敗 or 非大使クエリ → Claude 学習データで生成
        prompt = RECENT_ACTIVITIES_PROMPT.format(
            name_ja=info.get("full_name_ja", ""),
            name_native=name_native or info.get("full_name_ja", ""),
            nationality=info.get("nationality", ""),
            current_position=info.get("current_position", ""),
            search_language=search_language,
        )

    try:
        msg = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
        raw = re.sub(r"\s*```\s*$", "", raw)
        result = json.loads(raw)
        # ニュースURLを情報源に追記
        if news_url:
            result.setdefault("activity_sources", [])
            if news_url not in result["activity_sources"]:
                result["activity_sources"].append(news_url)
        return result
    except Exception as e:
        print(f"  [最近の活動] 取得エラー: {e}  スキップします。")
        return {}


# ── Word 文書を生成 ───────────────────────────────────────────────────────────
def generate_doc(info: dict, photo_path, personal: dict, sns: dict,
                 activities: dict, output_path: Path):
    doc = Document()

    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.0)

    # ── タイトル ─────────────────────────────────────────────────────────────
    p_title = para(doc, "添付１", size=14, bold=True,
                   align=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)
    bdr = OxmlElement("w:bdr")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), "4")
    bdr.set(qn("w:space"), "0")
    bdr.set(qn("w:color"), "000000")
    p_title.runs[0]._r.get_or_add_rPr().append(bdr)

    # ── 基本情報 + 写真 ───────────────────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style   = "Table Grid"
    tbl.autofit = False
    left  = tbl.rows[0].cells[0]
    right = tbl.rows[0].cells[1]
    set_cell_width(left,  12.8)
    set_cell_width(right,  4.2)
    add_info_line(left, "氏　　名", info.get("full_name_ja", ""), first=True)
    add_info_line(left, "生年月日", info.get("birth_date", ""))
    add_info_line(left, "国　　籍", info.get("nationality", ""))
    add_info_line(left, "現　　職", info.get("current_position", ""))
    remove_table_borders(tbl)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(0)
    if photo_path and Path(photo_path).exists():
        rp.add_run().add_picture(str(photo_path), width=Cm(3.2))

    doc.add_paragraph()

    # ── 学歴 ─────────────────────────────────────────────────────────────────
    add_section_header(doc, "学　　　　　歴")
    edu_list = info.get("education", [])
    if edu_list:
        edu = doc.add_table(rows=len(edu_list), cols=2)
        edu.style   = "Table Grid"
        edu.autofit = False
        for i, e in enumerate(edu_list):
            set_cell_font(edu.rows[i].cells[0], e.get("period", ""))
            set_cell_font(edu.rows[i].cells[1], e.get("description", ""))
            set_cell_width(edu.rows[i].cells[0], 4.7)
            set_cell_width(edu.rows[i].cells[1], 11.8)
        remove_table_borders(edu)
    doc.add_paragraph()

    # ── 主な職歴 ─────────────────────────────────────────────────────────────
    add_section_header(doc, "主　な　職　歴")
    career_list = info.get("career", [])
    if career_list:
        ctbl = doc.add_table(rows=len(career_list), cols=2)
        ctbl.style   = "Table Grid"
        ctbl.autofit = False
        for i, c in enumerate(career_list):
            set_cell_font(ctbl.rows[i].cells[0], c.get("year", ""))
            set_cell_font(ctbl.rows[i].cells[1], c.get("description", ""))
            set_cell_width(ctbl.rows[i].cells[0], 4.7)
            set_cell_width(ctbl.rows[i].cells[1], 11.8)
        remove_table_borders(ctbl)
    doc.add_paragraph()

    # ── パーソナル情報 ────────────────────────────────────────────────────────
    has_personal = any(personal.get(k) for k in
                       ("family", "hobbies", "leisure", "pets", "personal_notes"))
    if has_personal:
        add_section_header(doc, "パ　ー　ソ　ナ　ル　情　報　（公開情報）")

        def _join(lst):
            if not lst:
                return "不明"
            return "　/　".join(str(x) for x in lst if x and x != "不明")

        rows = []
        if personal.get("family"):
            rows.append(("家 族 構 成", _join(personal["family"])))
        if personal.get("hobbies"):
            rows.append(("趣　　　味", _join(personal["hobbies"])))
        if personal.get("leisure"):
            rows.append(("余暇の過ごし方", _join(personal["leisure"])))
        if personal.get("pets"):
            rows.append(("ペ　ッ　ト", _join(personal["pets"])))
        if personal.get("personal_notes"):
            notes = [n for n in personal["personal_notes"] if n and n != "不明"]
            if notes:
                rows.append(("そ の 他", "　/　".join(notes)))

        add_kv_table(doc, rows)

        # 情報源
        sources = [s for s in personal.get("official_sources", [])
                   if s and s != "不明"]
        if sources:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(4)
            sp.paragraph_format.space_after  = Pt(2)
            r = sp.add_run("【情報源】 " + "　/　".join(sources))
            r.font.size  = Pt(9)
            r.font.name  = "MS明朝"
            r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            _set_east_asian_font(r, "MS明朝")
        doc.add_paragraph()

    # ── SNS・評判情報 ─────────────────────────────────────────────────────────
    has_sns = any(sns.get(k) for k in
                  ("x_twitter", "instagram", "other_sns",
                   "media_reputation", "notable_quotes"))
    if has_sns:
        add_section_header(doc, "SNS・評　判　情　報　（第三者の声）")

        add_sns_block(doc, "X（旧Twitter）",  sns.get("x_twitter", []))
        add_sns_block(doc, "Instagram",        sns.get("instagram", []))
        add_sns_block(doc, "その他SNS",        sns.get("other_sns", []))
        add_sns_block(doc, "メディア・評判",   sns.get("media_reputation", []))

        quotes = [q for q in sns.get("notable_quotes", []) if q and q != "不明"]
        if quotes:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run("［注目発言・引用］")
            r.font.size = Pt(F)
            r.font.bold = True
            r.font.name = "MS明朝"
            r.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
            _set_east_asian_font(r, "MS明朝")
            for q in quotes:
                qp = doc.add_paragraph()
                qp.paragraph_format.space_before = Pt(0)
                qp.paragraph_format.space_after  = Pt(4)
                qr = qp.add_run(f"「{q}」")
                qr.font.size  = Pt(F)
                qr.font.name  = "MS明朝"
                qr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                _set_east_asian_font(qr, "MS明朝")
        doc.add_paragraph()

    # ── 最近の活動 ────────────────────────────────────────────────────────────
    act_list = activities.get("recent_activities", [])
    if act_list:
        add_section_header(doc, "最　近　の　主　な　活　動")
        atbl = doc.add_table(rows=len(act_list), cols=3)
        atbl.style   = "Table Grid"
        atbl.autofit = False
        for i, a in enumerate(act_list):
            date_str  = a.get("date", "")
            title_str = a.get("title", "")
            detail    = a.get("purpose", "") + (
                ("　" + a.get("result", "")) if a.get("result") else "")
            set_cell_font(atbl.rows[i].cells[0], date_str)
            set_cell_font(atbl.rows[i].cells[1], title_str, bold=True)
            set_cell_font(atbl.rows[i].cells[2], detail)
            set_cell_width(atbl.rows[i].cells[0], 3.0)
            set_cell_width(atbl.rows[i].cells[1], 4.5)
            set_cell_width(atbl.rows[i].cells[2], 9.0)
        remove_table_borders(atbl)

        act_sources = [s for s in activities.get("activity_sources", [])
                       if s and s != "不明"]
        if act_sources:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(4)
            sp.paragraph_format.space_after  = Pt(2)
            r = sp.add_run("【情報源】 " + "　/　".join(act_sources))
            r.font.size  = Pt(9)
            r.font.name  = "MS明朝"
            r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            _set_east_asian_font(r, "MS明朝")
        doc.add_paragraph()

    para(doc, "以　　上", size=F, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)

    doc.save(str(output_path))


# ── メイン ────────────────────────────────────────────────────────────────────
def main():
    CV_DIR.mkdir(parents=True, exist_ok=True)

    print("=" * 50)
    print("  略歴書ジェネレーター（拡張版）")
    print("=" * 50)
    try:
        name_input = input(
            "\n誰の略歴書を作成しますか？\n"
            "（例: マクロン仏大統領、Franck Leroy、習近平）\n> "
        ).strip()
        nationality_hint = input(
            "国籍・国名（省略可、例: フランス）\n> "
        ).strip()
        position_hint = input(
            "現職・役職（省略可、例: グラン・エスト州知事）\n> "
        ).strip()
    except (EOFError, KeyboardInterrupt):
        print("\nキャンセルされました。")
        return

    if not name_input:
        sys.exit("名前が入力されませんでした。")

    sys.stdout.flush()

    if not CLAUDE_API_KEY:
        sys.exit("ANTHROPIC_API_KEY 環境変数が設定されていません。")
    client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)

    # 0a. 大使クエリの場合、外務省 data.html から現在の大使名を取得して上書き
    print(f"\n  [大使判定] クエリを分析中: 「{name_input}」")
    resolved_name, ambassador_country, ambassador_type = detect_ambassador_query(name_input, client)
    if resolved_name != name_input:
        print(f"\n  [大使判定] クエリを解決: 「{name_input}」→「{resolved_name}」")
        # 大使国名が判明した場合は nationality_hint に反映
        if ambassador_country and not nationality_hint:
            nationality_hint = ambassador_country

    # 0b. 名前解決：日本語名・現地語名・検索言語を特定
    print(f"\n  [名前解決] 現地語名・検索言語を特定中: 「{resolved_name}」")
    name_resolved = resolve_name_multilingual(
        resolved_name, nationality_hint, position_hint, client)
    name_ja       = name_resolved.get("full_name_ja", resolved_name)
    name_native   = name_resolved.get("full_name_native", resolved_name)
    search_language = name_resolved.get("search_language", "日本語")
    # resolve で得た nationality/current_position をヒントとして保持
    if not nationality_hint:
        nationality_hint = name_resolved.get("nationality", "")
    if not position_hint:
        position_hint = name_resolved.get("current_position", "")
    print(f"  日本語名: {name_ja}")
    print(f"  現地語名: {name_native}")
    print(f"  検索言語: {search_language}")

    print(f"\n  [Claude] 情報を問い合わせ中: {name_ja} ({name_native}) ...")
    sys.stdout.flush()

    # 1. 人物基本情報を Claude で取得
    try:
        info = fetch_person_info(
            name_ja, client,
            name_native=name_native,
            nationality=nationality_hint,
            current_position=position_hint,
            search_language=search_language,
        )
    except Exception as e:
        sys.exit(f"情報取得に失敗しました: {e}")

    full_name_ja = info.get("full_name_ja", name_ja)
    print(f"\n  氏名: {full_name_ja}")
    print(f"  国籍: {info.get('nationality', '?')}")
    print(f"  現職: {info.get('current_position', '?')}")

    # 2. DB フォルダーから情報を取得して補強
    print("\n  [DB] 関連フォルダーを判定中...")
    db_folder = detect_db_folder(info, client)
    if db_folder:
        print("  [DB] ファイルを読み込み中...")
        files_block = build_files_block(db_folder)
        if files_block:
            file_count = files_block.count("=== ")
            print(f"  [DB] {file_count} 件のファイルから情報を取得。Claude で分析補強中...")
            info = enrich_with_db(info, files_block, client, name_native=name_native)
            print("  [DB] 補強完了。")
        else:
            print("  [DB] 読み込めるファイルがありませんでした。")
    else:
        print("  [DB] 該当フォルダーなし。スキップします。")

    # 3. パーソナル情報を取得
    print("\n  [パーソナル情報] 公式情報・懇談話題を取得中...")
    personal = fetch_personal_info(info, client,
                                   name_native=name_native,
                                   search_language=search_language)
    if personal:
        print("  [パーソナル情報] 取得完了。")
    else:
        print("  [パーソナル情報] 取得できませんでした。")

    # 4. SNS・評判情報を取得
    print("\n  [SNS情報] X・Instagram等の評判を取得中...")
    sns = fetch_sns_info(info, client,
                         name_native=name_native,
                         search_language=search_language)
    if sns:
        print("  [SNS情報] 取得完了。")
    else:
        print("  [SNS情報] 取得できませんでした。")

    # 4a. 大使クエリの場合、大使館ページURLを1回だけ取得（写真・活動で共用）
    embassy_pages = {}
    if ambassador_country:
        print(f"\n  [大使館] {ambassador_country}大使館ページを解析中...")
        embassy_pages = get_embassy_pages(
            ambassador_type, ambassador_country, full_name_ja, client)

    # 4b. 最近の活動を取得
    print("\n  [最近の活動] 大使館HP・ニュース等から活動を取得中...")
    activities = fetch_recent_activities(
        info, client, ambassador_country, ambassador_type, embassy_pages,
        name_native=name_native, search_language=search_language)
    if activities.get("recent_activities"):
        count = len(activities["recent_activities"])
        print(f"  [最近の活動] {count} 件取得完了。")
    else:
        print("  [最近の活動] 取得できませんでした。")

    # 5. 写真を取得
    safe_name  = re.sub(r'[\\/:*?"<>|\s]', "_", full_name_ja)
    photo_dir  = CV_DIR / "photo"
    photo_dir.mkdir(parents=True, exist_ok=True)
    photo_path = photo_dir / f"{safe_name}.jpg"

    has_photo = False

    # 公式サイトから優先取得（Claude が URL を提案）
    print("\n  [写真] 公式サイトから取得を試みています...")
    has_photo = fetch_official_photo(info, client, photo_path, name_native=name_native)

    # 公式サイト失敗時：大使クエリなら大使館サイトをスクレイピング
    if not has_photo and ambassador_country:
        print("\n  [写真] 大使館公式サイトから取得を試みています...")
        has_photo = fetch_ambassador_photo(
            info, ambassador_type, ambassador_country, client, photo_path,
            embassy_pages=embassy_pages)

    # Wikipedia フォールバック
    if not has_photo:
        print("  [写真] Wikipedia から取得を試みています...")
        has_photo = fetch_wikipedia_photo(info.get("wikipedia_title_en", ""), photo_path)

    if not has_photo:
        photo_path = None

    # 6. Word 文書を生成
    date_tag    = datetime.now().strftime("%Y%m%d")
    output_path = CV_DIR / f"{safe_name}_略歴書Plus_{date_tag}.docx"

    print("\n  [Word] 文書を生成中...")
    try:
        generate_doc(info, photo_path, personal, sns, activities, output_path)
        print(f"\n保存完了: {output_path}")
    except Exception as e:
        sys.exit(f"文書生成に失敗しました: {e}")


if __name__ == "__main__":
    main()
