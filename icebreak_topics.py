"""
icebreak_topics.py
==================
make_cv_plus.py の情報収集基盤を流用し、
懇談アイスブレイク用のトピック案（3〜5件）を選定して
Word 文書で出力するツール。

出力先: C:\\Users\\shondo\\Desktop\\agent_project\\cv
依存:   pip install anthropic python-docx pillow
"""

import io
import json
import re
import sys
import urllib.request
import urllib.parse
from pathlib import Path
from datetime import datetime

if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
    sys.stdin  = io.TextIOWrapper(sys.stdin.buffer,  encoding="utf-8", errors="replace")

try:
    import anthropic
except ImportError:
    sys.exit("anthropic が未インストールです。pip install anthropic を実行してください。")

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("python-docx が未インストールです。pip install python-docx を実行してください。")

try:
    import fitz
except ImportError:
    fitz = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

# ── 設定 ─────────────────────────────────────────────────────────────────────
CLAUDE_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
CLAUDE_MODEL   = "claude-sonnet-4-6"
CV_DIR         = Path(r"C:\Users\shondo\Desktop\agent_project\cv")
DB_BASE        = Path(r"C:\Users\shondo\Desktop\agent_project\db")
F              = 12
HEADERS        = {"User-Agent": "Mozilla/5.0 (compatible; CVMaker/1.0)"}
MOFA_HEADERS   = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept": (
        "text/html,application/xhtml+xml,application/xml;q=0.9,"
        "image/avif,image/webp,*/*;q=0.8"
    ),
    "Accept-Language": "ja,en-US;q=0.9,en;q=0.8",
    "Accept-Encoding": "gzip, deflate, br",
    "Referer": "https://www.mofa.go.jp/",
    "Connection": "keep-alive",
}
GENERIC_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "ja,en-US;q=0.9,en;q=0.8",
    "Accept-Encoding": "gzip, deflate, br",
    "Connection": "keep-alive",
}
CHARS_PER_FILE = 4000

DB_FOLDERS = {
    "asean_oceania": "ASEAN諸国・オセアニア（タイ、ベトナム、インドネシア、シンガポール、豪州など）",
    "china":         "中国（習近平政権、中国共産党、人民元、香港など）",
    "europe":        "欧州（EU、英独仏伊、ウクライナ、ロシア、NATO欧州側など）",
    "geopolitics":   "地政学・地経学全般（半導体、サプライチェーン、鉱物資源など）",
    "india_mena":    "インド・中東・アフリカ（サウジ、イスラエル、イラン、UAE、ケニアなど）",
    "japan":         "日本（日本政府、日本企業、日本の外交・経済政策など）",
    "minerals":      "鉱物資源（レアアース、ニッケル、リチウム、黒鉛など）",
    "north america": "北米（米国、トランプ政権、カナダ、メキシコ、FRBなど）",
    "south america": "南米（ブラジル、アルゼンチン、ベネズエラなど）",
    "taiwan":        "台湾（台湾政府、台湾海峡、TSMC、台湾の安全保障など）",
    "trade_tariff_ai": "通商・関税・AI（WTO、FTA、人権、環境、AI規制など）",
    "others":        "上記以外",
}

# ── Claude プロンプト（基本情報） ─────────────────────────────────────────────
INFO_PROMPT = """\
あなたは政治家・外交官・有識者の情報を整理する専門家です。
以下の人物について、日本の外交官が面談準備に使う略歴書用の情報を
JSON 形式で提供してください。

【対象人物】
{name}

【出力ルール】
- JSON のみを出力（前置き・説明・コードブロック記号は一切不要）
- 各フィールドの文字数は必ず守ること（Word 文書の 1 行に収める制約）
- 情報が不確かな場合は "不明" と記載

{{
  "full_name_ja": "日本語・カタカナ表記の氏名",
  "birth_date": "生年月日（例: 1957年 4月 22日）",
  "nationality": "国籍（例: ポーランド共和国）",
  "current_position": "現職（30文字以内）",
  "education": [
    {{"period": "在籍期間（例: 1975年 ～ 1980年）",
      "description": "学校名・学部（25文字以内）"}}
  ],
  "career": [
    {{"year": "年（例: 2007年）",
      "description": "役職・出来事（25文字以内、現職・功績に関係するものに絞る、最大 8 件）"}}
  ],
  "interests": ["関心事項・主要政策（各28文字以内、5〜7件）"],
  "todo": ["面談で歓迎される話題・姿勢（各28文字以内、5〜7件）"],
  "not_todo": ["面談で避けるべき言動（各28文字以内、5〜7件）"],
  "wikipedia_title_en": "英語版 Wikipedia の記事タイトル（例: Donald Tusk）"
}}
"""

FOLDER_DETECT_PROMPT = """\
以下の人物が「最も関係する地域・テーマ」に対応するフォルダーを1つだけ選んでください。

【人物情報】
氏名: {name}
国籍: {nationality}
現職: {current_position}

【選択肢】
{folder_list}

回答はフォルダー名のみ1行で返してください（説明不要）。
"""

ENRICH_PROMPT = """\
あなたは外交官向け面談準備資料の専門家です。
以下の人物の略歴書草案に、データベースファイルの最新情報を反映して
interests・todo・not_todo を更新してください。

【対象人物】
氏名: {name}（{nationality}、{current_position}）

【データベースファイルの内容】
{files_block}

【現在の草案（interests / todo / not_todo）】
{current_sections}

【指示】
- データベースの情報を踏まえ、より具体的・最新の内容に更新する
- 草案の内容も活かしつつ、DB 情報で補強・修正する
- 各項目は28文字以内
- interests: 5〜7件、todo: 5〜7件、not_todo: 5〜7件
- JSON のみを出力（前置き不要）。形式:
{{"interests": [...], "todo": [...], "not_todo": [...]}}
"""

PERSONAL_INFO_PROMPT = """\
あなたは政治家・外交官・有識者に関する公開情報を調査する専門家です。
以下の人物について、公式ウェブサイト・公的プロフィール・インタビュー記事など
「公的・公開情報」をもとに、パーソナルな情報と懇談に役立つ話題を提供してください。

【対象人物】
氏名: {name}
国籍: {nationality}
現職: {current_position}

【出力ルール】
- JSON のみを出力（前置き・説明・コードブロック記号は一切不要）
- 確認できた情報のみ記載。不明な場合は "不明" または空リスト []
- 推測・伝聞は「〜とされる」「〜と伝えられる」と明記
- 各文字列は40文字以内

{{
  "family": [
    "家族構成・配偶者・子供など（公開情報のみ、各40文字以内）"
  ],
  "hobbies": [
    "趣味・関心（例: ゴルフ、読書、クラシック音楽）（各30文字以内）"
  ],
  "leisure": [
    "余暇の過ごし方（例: テニス、ハイキング、料理）（各30文字以内）"
  ],
  "pets": [
    "ペット情報（例: 犬2匹（ゴールデンレトリバー）など、不明な場合は 不明）"
  ],
  "personal_notes": [
    "その他パーソナルな公開情報（出身地、食の好み、信条など）（各40文字以内）"
  ],
  "conversation_topics": [
    "有識者懇談で盛り上がる可能性のある話題（各40文字以内、5〜7件）"
  ],
  "official_sources": [
    "情報源となった公式サイト・インタビューの媒体名またはURL"
  ]
}}
"""

NEWS_PARSE_PROMPT = """\
以下のウェブサイトから取得したテキストを分析して、
{name}（{country}大使）の最近の主な活動・イベントを
新しい順（最新が先頭）で3〜5件まとめてください。

【サイトテキスト】
{text}

【出力ルール】
- JSON のみを出力（前置き・コードブロック不要）
- 活動の日付・内容が明確なものだけを抽出
- 不明・確認できない場合は空リスト []

{{
  "recent_activities": [
    {{
      "date": "活動日・期間（例: 2025年5月）",
      "title": "活動名（25文字以内）",
      "purpose": "目的・背景（50文字以内）",
      "result": "成果・内容（50文字以内）"
    }}
  ],
  "activity_sources": ["情報源URL"]
}}
"""

RECENT_ACTIVITIES_PROMPT = """\
あなたは外交官・政治家・有識者の公開情報を調査する専門家です。
以下の人物について、大使館公式サイト・外務省・ニュース記事・Google検索結果をもとに
「最近の主な活動」を時系列で調査してください。

【対象人物】
氏名: {name}
国籍: {nationality}
現職: {current_position}

【出力ルール】
- JSON のみを出力（前置き・コードブロック不要）
- 直近1〜2年の活動を優先し、新しい順（最新が先頭）で3〜5件
- 確認できた事実のみ記載。推測には「〜とみられる」と明記
- 不明・確認できない場合は空リスト []
- title は25文字以内、purpose・result は各50文字以内

{{
  "recent_activities": [
    {{
      "date": "活動日・期間（例: 2024年11月）",
      "title": "活動名・イベント名（25文字以内）",
      "purpose": "目的・背景（50文字以内）",
      "result": "成果・結果（50文字以内）"
    }}
  ],
  "activity_sources": [
    "情報源となったサイトURL または媒体名（確認できた範囲で）"
  ]
}}
"""

AMBASSADOR_DETECT_PROMPT = """\
以下の入力文字列が「大使または外交官」に関する調査依頼かどうかを判断してください。

【入力】{query}

【判断基準】
- 「チェコ大使」「駐日フランス大使」「特命全権大使」などが含まれる場合は大使関連
- 人名が入力された場合、その人物が大使・外交官として著名であれば大使関連
- それ以外（政治家・経済人・研究者など）は大使関連でない

【ambassador_type の定義】
  - "駐日大使": 相手国が日本に派遣している大使（例: 駐日チェコ大使）
  - "駐外大使": 日本が相手国に派遣している大使（例: 駐チェコ日本大使）
  - "不明": 種別が判断できない場合

【出力ルール】
- JSON のみを出力（前置き・コードブロック不要）

{{
  "is_ambassador": true または false,
  "ambassador_type": "駐日大使" または "駐外大使" または "不明",
  "country_name": "国名（例: チェコ、フランス、アメリカ）。不明な場合は空文字",
  "mofa_page_url": "外務省の当該国ページURL。不明な場合は空文字"
}}
"""

EMBASSY_BASE_URL_PROMPT = """\
以下の大使館の公式ウェブサイトのトップページURLを1件返してください。

種別: {ambassador_type}
相手国: {country}
大使名（日本語）: {name}

URLのみ1行で返してください（説明不要）。
"""

# ── 懇談アイスブレイクトピック選定プロンプト ─────────────────────────────────
ICEBREAK_TOPICS_PROMPT = """\
あなたは外交・儀礼の専門家です。
以下に示す人物に関する情報をもとに、懇談の場でアイスブレイクとして使える
「相手が喜ぶ・話したい」トピックを3〜5件選んでください。

【対象人物】
氏名: {name}
国籍: {nationality}
現職: {current_position}

【収集済み情報】
■ 最近の活動（新しい順）:
{recent_activities}

■ 関心事項・主要政策:
{interests}

■ 歓迎される話題・姿勢（To Do）:
{todo}

■ 避けるべき言動（Not to Do）:
{not_todo}

■ 懇談に役立つ話題（既出案）:
{conversation_topics}

■ パーソナル情報:
{personal_info}

【選定基準】
1. 相手が「喜ぶ」「話したがる」トピックを優先する
2. 意見の対立が起きにくい話題を選ぶ
3. ペット・趣味・余暇・スポーツ・音楽・食文化・出身地・最近の活動など
   文化的・個人的なテーマを積極的に採用する
4. 政治対立・ビジネス交渉・外交上の難題は原則として避ける
5. Not to Do に含まれるテーマは絶対に選ばない

【出力ルール】
- JSON のみを出力（前置き・説明・コードブロック記号は一切不要）
- 3〜5件を選定する

{{
  "topics": [
    {{
      "title": "トピックのタイトル（20文字以内）",
      "what_to_say": "具体的な話題の振り方・切り出し例文（80文字以内）",
      "why": "このトピックを選んだ理由・相手が喜ぶと考える根拠（100文字以内）",
      "background": "話題の背景・文脈・事前に知っておくと役立つ情報（150文字以内）"
    }}
  ]
}}
"""


# ── XML / フォントヘルパー ─────────────────────────────────────────────────────
def _set_east_asian_font(run, font_name: str):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "FFFFFF")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _set_table_left_align(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn("w:jc")):
        tblPr.remove(existing)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "left")
    tblPr.append(jc)


def set_cell_width(cell, width_cm: float):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")


def set_cell_font(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(F)
    run.font.bold = bold
    run.font.name = "MS明朝"
    _set_east_asian_font(run, "MS明朝")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


# ── 文書生成ヘルパー ──────────────────────────────────────────────────────────
def para(doc, text: str = "", size: int = F, bold: bool = False,
         align=WD_ALIGN_PARAGRAPH.LEFT, color=None,
         space_after: int = 3, underline: bool = False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.size      = Pt(size)
        run.font.bold      = bold
        run.font.underline = underline
        run.font.name      = "MS明朝"
        _set_east_asian_font(run, "MS明朝")
        if color:
            run.font.color.rgb = color
    return p


def add_section_header(doc, text: str):
    para(doc, text, size=F, bold=True, space_after=4, underline=True)


def _add_hr(doc):
    """段落の下枠線で水平区切り線を表現する。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_topic_table(doc, rows_data: list):
    """トピック詳細を2列テーブルで表示する（枠線あり）。"""
    if not rows_data:
        return
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style   = "Table Grid"
    tbl.autofit = False
    for i, (label, value) in enumerate(rows_data):
        set_cell_font(tbl.rows[i].cells[0], label, bold=True)
        set_cell_font(tbl.rows[i].cells[1], value)
        set_cell_width(tbl.rows[i].cells[0], 3.2)
        set_cell_width(tbl.rows[i].cells[1], 13.3)


# ── Claude で人物情報を取得 ────────────────────────────────────────────────────
def fetch_person_info(name: str, client) -> dict:
    msg = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=2000,
        messages=[{"role": "user", "content": INFO_PROMPT.format(name=name)}],
    )
    raw = msg.content[0].text.strip().lstrip("﻿")
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s*```\s*$", "", raw)
    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"  JSON パースエラー: {e}")
        print(f"  出力の先頭 300 文字:\n{raw[:300]}")
        raise


# ── DB フォルダーの判定 ────────────────────────────────────────────────────────
def detect_db_folder(info: dict, client) -> Path | None:
    available = {k: v for k, v in DB_FOLDERS.items()
                 if (DB_BASE / k).is_dir()}
    if not available:
        return None
    folder_list = "\n".join(f"  {k}: {v}" for k, v in available.items())
    prompt = FOLDER_DETECT_PROMPT.format(
        name=info.get("full_name_ja", ""),
        nationality=info.get("nationality", ""),
        current_position=info.get("current_position", ""),
        folder_list=folder_list,
    )
    msg = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=30,
        messages=[{"role": "user", "content": prompt}],
    )
    chosen = msg.content[0].text.strip().lower().strip('"').strip("'")
    for key in available:
        if key in chosen:
            folder = DB_BASE / key
            print(f"  [DB] 関連フォルダー: {folder}")
            return folder
    return None


# ── DB ファイルのテキスト抽出 ─────────────────────────────────────────────────
def _extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            if pdfplumber:
                with pdfplumber.open(str(path)) as pdf:
                    return "\n".join(p.extract_text() or "" for p in pdf.pages)
            if fitz:
                doc = fitz.open(str(path))
                return "\n".join(page.get_text() for page in doc)
            return ""
        if ext == ".docx":
            doc = Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        if ext in (".txt", ".md"):
            return path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        pass
    return ""


def build_files_block(folder: Path) -> str:
    supported = {".pdf", ".docx", ".txt", ".md"}
    files = sorted(folder.rglob("*"))
    blocks = []
    for f in files:
        if f.suffix.lower() not in supported:
            continue
        text = _extract_text(f)
        if not text.strip():
            continue
        snippet = text[:CHARS_PER_FILE]
        if len(text) > CHARS_PER_FILE:
            snippet += "\n[以下省略]"
        blocks.append(f"=== {f.name} ===\n{snippet}")
        if len(blocks) >= 15:
            break
    return "\n\n".join(blocks)


# ── DB 情報で interests / todo / not_todo を補強 ──────────────────────────────
def enrich_with_db(info: dict, files_block: str, client) -> dict:
    current_sections = json.dumps({
        "interests": info.get("interests", []),
        "todo":      info.get("todo", []),
        "not_todo":  info.get("not_todo", []),
    }, ensure_ascii=False, indent=2)
    prompt = ENRICH_PROMPT.format(
        name=info.get("full_name_ja", ""),
        nationality=info.get("nationality", ""),
        current_position=info.get("current_position", ""),
        files_block=files_block,
        current_sections=current_sections,
    )
    msg = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = msg.content[0].text.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s*```\s*$", "", raw)
    try:
        enriched = json.loads(raw)
        for key in ("interests", "todo", "not_todo"):
            if key in enriched and isinstance(enriched[key], list):
                info[key] = enriched[key]
        return info
    except json.JSONDecodeError as e:
        print(f"  [DB補強] JSON パースエラー: {e}  草案のまま使用します。")
        return info


# ── パーソナル情報を取得 ───────────────────────────────────────────────────────
def fetch_personal_info(info: dict, client) -> dict:
    prompt = PERSONAL_INFO_PROMPT.format(
        name=info.get("full_name_ja", ""),
        nationality=info.get("nationality", ""),
        current_position=info.get("current_position", ""),
    )
    try:
        msg = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=1500,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
        raw = re.sub(r"\s*```\s*$", "", raw)
        return json.loads(raw)
    except Exception as e:
        print(f"  [パーソナル情報] 取得エラー: {e}  スキップします。")
        return {}


# ── 外務省・大使館スクレイピング ──────────────────────────────────────────────
def _fetch_mofa_html(url: str) -> str:
    import gzip as _gzip
    req = urllib.request.Request(url, headers=MOFA_HEADERS)
    with urllib.request.urlopen(req, timeout=15) as resp:
        raw = resp.read()
    try:
        return _gzip.decompress(raw).decode("utf-8", errors="replace")
    except Exception:
        return raw.decode("utf-8", errors="replace")


def _fetch_html_generic(url: str) -> str:
    import gzip as _gzip
    req = urllib.request.Request(url, headers=GENERIC_HEADERS)
    with urllib.request.urlopen(req, timeout=15) as resp:
        raw     = resp.read()
        charset = resp.headers.get_content_charset() or "utf-8"
    try:
        decoded = _gzip.decompress(raw)
    except Exception:
        decoded = raw
    for enc in (charset, "utf-8", "latin-1"):
        try:
            return decoded.decode(enc, errors="replace")
        except Exception:
            pass
    return decoded.decode("utf-8", errors="replace")


def _abs_url(src: str, base_url: str) -> str:
    if src.startswith("http"):
        return src
    if src.startswith("//"):
        return "https:" + src
    p = urllib.parse.urlparse(base_url)
    if src.startswith("/"):
        return f"{p.scheme}://{p.netloc}{src}"
    return base_url.rsplit("/", 1)[0] + "/" + src


def _extract_hrefs(html: str) -> list:
    return re.findall(r'href=["\']([^"\']+)["\']', html, re.IGNORECASE)


def _find_ambassador_page(hrefs: list, base_url: str) -> str:
    bad = ("svg","png","jpg","css","js","ico","pdf","#","?","mail","tel")
    AMB_WORDS = (
        "ambassador", "velvyslanec", "ambassadeur", "botschafter",
        "embajador", "ambasciatore", "ambasador", "ambassadör",
    )
    for lk in hrefs:
        lk_low = lk.lower()
        if any(w in lk_low for w in AMB_WORDS):
            if not any(lk_low.endswith(x) for x in bad):
                return _abs_url(lk, base_url)
    return ""


def _find_news_page(hrefs: list, base_url: str) -> str:
    bad = ("svg","png","jpg","css","js","ico","pdf","#","?","mail","tel")
    keywords = ("bilateral","news","press","activity","event","relations","aktualne","aktual")
    candidates = []
    for lk in hrefs:
        if any(k in lk.lower() for k in keywords):
            if not any(lk.lower().endswith(x) for x in bad):
                candidates.append(lk)
    if not candidates:
        return ""
    candidates.sort(key=lambda x: (0 if "index" in x.lower() else 1, len(x)))
    return _abs_url(candidates[0], base_url)


def get_embassy_pages(ambassador_type: str, country: str, name: str, client) -> dict:
    prompt = EMBASSY_BASE_URL_PROMPT.format(
        ambassador_type=ambassador_type or "大使",
        country=country,
        name=name,
    )
    try:
        msg = client.messages.create(
            model=CLAUDE_MODEL, max_tokens=100,
            messages=[{"role": "user", "content": prompt}],
        )
        base_url = msg.content[0].text.strip()
        if not base_url.startswith("http"):
            return {}
    except Exception as e:
        print(f"  [大使館] ベースURL取得エラー: {e}")
        return {}

    print(f"  [大使館] サイトを解析中: {base_url}")
    try:
        html_top = _fetch_html_generic(base_url)
    except Exception as e:
        print(f"  [大使館] サイト取得失敗: {e}")
        return {"embassy_base": base_url}

    hrefs_top = _extract_hrefs(html_top)
    result = {"embassy_base": base_url}

    amb_page = _find_ambassador_page(hrefs_top, base_url)
    if not amb_page:
        sub_keywords = (
            "embassy", "about", "staff", "mission", "consulate", "diplomatic",
            "velvyslanectvi", "o_velvyslanectvi",
            "embajada", "botschaft", "ambasada", "ambassade",
        )
        sub_candidates = []
        for lk in hrefs_top:
            if any(k in lk.lower() for k in sub_keywords):
                if not any(lk.lower().endswith(x) for x in ("svg","png","jpg","css","js","#","?")):
                    sub_candidates.append(_abs_url(lk, base_url))
        seen = set()
        sub_unique = []
        for u in sorted(sub_candidates, key=lambda x: (0 if "index" in x.lower() else 1, len(x))):
            if u not in seen:
                seen.add(u)
                sub_unique.append(u)
        for sub_url in sub_unique[:4]:
            print(f"  [大使館] サブページを確認: {sub_url}")
            try:
                html_sub = _fetch_html_generic(sub_url)
                hrefs_sub = _extract_hrefs(html_sub)
                amb_page = _find_ambassador_page(hrefs_sub, sub_url)
                if amb_page:
                    break
            except Exception:
                continue

    if amb_page:
        result["ambassador_page"] = amb_page
        print(f"  [大使館] 大使紹介ページ: {amb_page}")

    news_page = _find_news_page(hrefs_top, base_url)
    if news_page:
        result["news_page"] = news_page
        print(f"  [大使館] ニュースページ: {news_page}")

    return result


def scrape_mofa_ambassador_name(mofa_page_url: str, ambassador_type: str):
    if not mofa_page_url:
        return None
    data_url = re.sub(r'index\.html$', 'data.html', mofa_page_url)
    if data_url == mofa_page_url:
        data_url = mofa_page_url.rstrip("/") + "/data.html"
    print(f"  [MOFA大使] データページを取得中: {data_url}")
    try:
        html = _fetch_mofa_html(data_url)
    except Exception as e:
        print(f"  [MOFA大使] データページ取得失敗: {e}")
        return None
    idx = html.find("外交使節")
    if idx == -1:
        return None
    section = html[idx: idx + 800]
    items = re.findall(r'<li[^>]*>(.*?)</li>', section, re.DOTALL)
    for item in items:
        text = re.sub(r'<[^>]+>', '', item).strip()
        text = re.sub(r'^[（(]\d+[）)]\s*', '', text).strip()
        if not text:
            continue
        if ambassador_type == "駐日大使" and "駐日" in text:
            print(f"  [MOFA大使] 駐日大使を検出: {text}")
            return text
        elif ambassador_type == "駐外大使" and "駐日" not in text:
            print(f"  [MOFA大使] 駐外大使を検出: {text}")
            return text
    if ambassador_type == "不明" and items:
        text = re.sub(r'<[^>]+>', '', items[0]).strip()
        text = re.sub(r'^[（(]\d+[）)]\s*', '', text).strip()
        if text:
            return text
    return None


def detect_ambassador_query(query: str, client):
    prompt = AMBASSADOR_DETECT_PROMPT.format(query=query)
    try:
        msg = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=400,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
        raw = re.sub(r"\s*```\s*$", "", raw)
        result = json.loads(raw)
    except Exception as e:
        print(f"  [大使判定] エラー: {e}  スキップします。")
        return query, "", ""
    if not result.get("is_ambassador", False):
        return query, "", ""
    country  = result.get("country_name", "")
    mofa_url = result.get("mofa_page_url", "")
    amb_type = result.get("ambassador_type", "不明")
    print(f"  [大使判定] 大使クエリを検出: 国={country}, 種別={amb_type}")
    if not mofa_url:
        return query, country, amb_type
    ambassador_name = scrape_mofa_ambassador_name(mofa_url, amb_type)
    if ambassador_name:
        print(f"  [大使判定] 現在の大使名を確認: 「{ambassador_name}」")
        return ambassador_name, country, amb_type
    return query, country, amb_type


# ── 最近の活動を取得 ──────────────────────────────────────────────────────────
def fetch_recent_activities(info: dict, client,
                            ambassador_country: str = "",
                            ambassador_type: str = "",
                            embassy_pages: dict = None) -> dict:
    scraped_text = ""
    news_url     = ""
    if ambassador_country and embassy_pages:
        news_url = embassy_pages.get("news_page", "")
        if not news_url:
            base_url = embassy_pages.get("embassy_base", "")
            if base_url:
                try:
                    html_base = _fetch_html_generic(base_url)
                    for m in re.finditer(r'href=["\']([^"\'#?][^"\']*)["\']', html_base):
                        lk = m.group(1)
                        if any(k in lk.lower()
                               for k in ["bilateral","news","press","activity","event"]):
                            if not lk.lower().endswith(("svg","png","jpg","css","js")):
                                news_url = _abs_url(lk, base_url)
                                break
                except Exception:
                    pass
        if news_url and news_url.startswith("http"):
            print(f"  [最近の活動] ニュースページを取得中: {news_url}")
            try:
                html = _fetch_html_generic(news_url)
                text = re.sub(r'<[^>]+>', ' ', html)
                text = re.sub(r'&[a-z]+;', ' ', text)
                text = re.sub(r'\s+', ' ', text).strip()
                scraped_text = text[:5000]
            except Exception as e:
                print(f"  [最近の活動] ニュースページ取得失敗: {e}")
    if scraped_text:
        prompt = NEWS_PARSE_PROMPT.format(
            name=info.get("full_name_ja", ""),
            country=ambassador_country or info.get("nationality", ""),
            text=scraped_text,
        )
    else:
        prompt = RECENT_ACTIVITIES_PROMPT.format(
            name=info.get("full_name_ja", ""),
            nationality=info.get("nationality", ""),
            current_position=info.get("current_position", ""),
        )
    try:
        msg = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
        raw = re.sub(r"\s*```\s*$", "", raw)
        result = json.loads(raw)
        if news_url:
            result.setdefault("activity_sources", [])
            if news_url not in result["activity_sources"]:
                result["activity_sources"].append(news_url)
        return result
    except Exception as e:
        print(f"  [最近の活動] 取得エラー: {e}  スキップします。")
        return {}


# ── アイスブレイクトピックを選定 ─────────────────────────────────────────────
def select_icebreak_topics(info: dict, personal: dict,
                           activities: dict, client) -> list:
    """収集済み情報から懇談アイスブレイク用トピックを3〜5件選定する。"""

    acts = activities.get("recent_activities", [])
    recent_str = "\n".join(
        f"・{a.get('date', '')} {a.get('title', '')}：{a.get('purpose', '')} / {a.get('result', '')}"
        for a in acts
    ) if acts else "（情報なし）"

    interests_str      = "、".join(info.get("interests", [])) or "（情報なし）"
    todo_str           = "、".join(info.get("todo", []))      or "（情報なし）"
    not_todo_str       = "、".join(info.get("not_todo", []))  or "（情報なし）"
    conv_topics_str    = "、".join(personal.get("conversation_topics", [])) or "（情報なし）"

    personal_parts = []
    if personal.get("family"):
        personal_parts.append("家族: " + " / ".join(personal["family"]))
    if personal.get("hobbies"):
        personal_parts.append("趣味: " + " / ".join(personal["hobbies"]))
    if personal.get("leisure"):
        personal_parts.append("余暇: " + " / ".join(personal["leisure"]))
    pets = [p for p in personal.get("pets", []) if p and p != "不明"]
    if pets:
        personal_parts.append("ペット: " + " / ".join(pets))
    if personal.get("personal_notes"):
        personal_parts.append("その他: " + " / ".join(personal["personal_notes"]))
    personal_str = "\n".join(personal_parts) or "（情報なし）"

    prompt = ICEBREAK_TOPICS_PROMPT.format(
        name=info.get("full_name_ja", ""),
        nationality=info.get("nationality", ""),
        current_position=info.get("current_position", ""),
        recent_activities=recent_str,
        interests=interests_str,
        todo=todo_str,
        not_todo=not_todo_str,
        conversation_topics=conv_topics_str,
        personal_info=personal_str,
    )

    try:
        msg = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
        raw = re.sub(r"\s*```\s*$", "", raw)
        result = json.loads(raw)
        return result.get("topics", [])
    except Exception as e:
        print(f"  [トピック選定] エラー: {e}  スキップします。")
        return []


# ── Word 文書を生成 ───────────────────────────────────────────────────────────
def generate_icebreak_doc(name: str, topics: list, output_path: Path):
    doc = Document()

    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.0)

    # ── 「添付４」囲み線ボックス（左端） ────────────────────────────────────────
    att_tbl = doc.add_table(rows=1, cols=1)
    att_tbl.style   = "Table Grid"
    att_tbl.autofit = False
    att_cell = att_tbl.rows[0].cells[0]
    set_cell_width(att_cell, 2.0)
    att_cell.text = ""
    att_p = att_cell.paragraphs[0]
    att_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    att_run = att_p.add_run("添付４")
    att_run.font.size = Pt(F)
    att_run.font.bold = True
    att_run.font.name = "MS明朝"
    _set_east_asian_font(att_run, "MS明朝")
    att_p.paragraph_format.space_before = Pt(3)
    att_p.paragraph_format.space_after  = Pt(3)
    _set_table_left_align(att_tbl)

    doc.add_paragraph()

    # ── タイトル ─────────────────────────────────────────────────────────────
    para(doc, "懇　談　ト　ピ　ッ　ク　案", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    para(doc, f"（対象者: {name}）", size=F,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    _add_hr(doc)

    # ── 各トピック ────────────────────────────────────────────────────────────
    for i, topic in enumerate(topics, 1):
        title       = topic.get("title", "")
        what_to_say = topic.get("what_to_say", "")
        why         = topic.get("why", "")
        background  = topic.get("background", "")

        add_section_header(doc, f"【トピック{i}】　{title}")

        rows_data = []
        if what_to_say:
            rows_data.append(("話題の振り方", what_to_say))
        if why:
            rows_data.append(("選　定　理　由", why))
        if background:
            rows_data.append(("背景・補足情報", background))

        add_topic_table(doc, rows_data)
        doc.add_paragraph()

    # ── 以上 ──────────────────────────────────────────────────────────────────
    para(doc, "以　　上", size=F, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)

    doc.save(str(output_path))


# ── メイン ────────────────────────────────────────────────────────────────────
def main():
    CV_DIR.mkdir(parents=True, exist_ok=True)

    print("=" * 50)
    print("  懇談トピック案ジェネレーター")
    print("=" * 50)
    try:
        name = input(
            "\n誰との懇談トピックを作成しますか？\n"
            "（例: マクロン仏大統領、習近平、チェコ大使）\n> "
        ).strip()
    except (EOFError, KeyboardInterrupt):
        print("\nキャンセルされました。")
        return

    if not name:
        sys.exit("名前が入力されませんでした。")

    sys.stdout.flush()

    client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)

    # 0. 大使クエリの場合、外務省から現在の大使名を取得
    print(f"\n  [大使判定] クエリを分析中: 「{name}」")
    resolved_name, ambassador_country, ambassador_type = detect_ambassador_query(name, client)
    if resolved_name != name:
        print(f"\n  [大使判定] クエリを解決: 「{name}」→「{resolved_name}」")
        name = resolved_name

    print(f"\n  [Claude] 情報を問い合わせ中: {name} ...")
    sys.stdout.flush()

    # 1. 人物基本情報を取得
    try:
        info = fetch_person_info(name, client)
    except Exception as e:
        sys.exit(f"情報取得に失敗しました: {e}")

    full_name_ja = info.get("full_name_ja", name)
    print(f"\n  氏名: {full_name_ja}")
    print(f"  国籍: {info.get('nationality', '?')}")
    print(f"  現職: {info.get('current_position', '?')}")

    # 2. DB フォルダーから情報を取得して補強
    print("\n  [DB] 関連フォルダーを判定中...")
    db_folder = detect_db_folder(info, client)
    if db_folder:
        print("  [DB] ファイルを読み込み中...")
        files_block = build_files_block(db_folder)
        if files_block:
            file_count = files_block.count("=== ")
            print(f"  [DB] {file_count} 件のファイルから情報を取得。Claude で分析補強中...")
            info = enrich_with_db(info, files_block, client)
            print("  [DB] 補強完了。")
        else:
            print("  [DB] 読み込めるファイルがありませんでした。")
    else:
        print("  [DB] 該当フォルダーなし。スキップします。")

    # 3. パーソナル情報を取得
    print("\n  [パーソナル情報] 公式情報・懇談話題を取得中...")
    personal = fetch_personal_info(info, client)
    if personal:
        print("  [パーソナル情報] 取得完了。")
    else:
        print("  [パーソナル情報] 取得できませんでした。")

    # 4. 大使クエリの場合、大使館ページを解析
    embassy_pages = {}
    if ambassador_country:
        print(f"\n  [大使館] {ambassador_country}大使館ページを解析中...")
        embassy_pages = get_embassy_pages(
            ambassador_type, ambassador_country, full_name_ja, client)

    # 5. 最近の活動を取得
    print("\n  [最近の活動] 大使館HP・ニュース等から活動を取得中...")
    activities = fetch_recent_activities(
        info, client, ambassador_country, ambassador_type, embassy_pages)
    count = len(activities.get("recent_activities", []))
    if count:
        print(f"  [最近の活動] {count} 件取得完了。")
    else:
        print("  [最近の活動] 取得できませんでした。")

    # 6. アイスブレイクトピックを選定
    print("\n  [トピック選定] 懇談トピックを選定中...")
    topics = select_icebreak_topics(info, personal, activities, client)
    if topics:
        print(f"  [トピック選定] {len(topics)} 件のトピックを選定しました。")
        for i, t in enumerate(topics, 1):
            print(f"    {i}. {t.get('title', '')}")
    else:
        sys.exit("  トピックを選定できませんでした。処理を中止します。")

    # 7. Word 文書を生成
    safe_name   = re.sub(r'[\\/:*?"<>|\s]', "_", full_name_ja)
    date_tag    = datetime.now().strftime("%Y%m%d")
    output_path = CV_DIR / f"{safe_name}_懇談トピック案_{date_tag}.docx"

    print("\n  [Word] 文書を生成中...")
    try:
        generate_icebreak_doc(full_name_ja, topics, output_path)
        print(f"\n保存完了: {output_path}")
    except Exception as e:
        sys.exit(f"文書生成に失敗しました: {e}")


if __name__ == "__main__":
    main()
