"""
make_cv.py
==========
人物名を入力すると Claude API で略歴・関心事項を調査し、
A4 MS Word 略歴書（_make_tusk_doc2.py と同じ体裁）を生成するツール。

出力先: C:\\Users\\shondo\\Desktop\\agent_project\\cv

依存:
  pip install anthropic python-docx pillow
"""

import io
import json
import re
import sys
import urllib.request
import urllib.parse
from io import BytesIO
from pathlib import Path
from datetime import datetime

# Windows コンソールでの文字化けを防ぐ
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
    sys.stdin  = io.TextIOWrapper(sys.stdin.buffer,  encoding="utf-8", errors="replace")

try:
    import anthropic
except ImportError:
    sys.exit("anthropic が未インストールです。pip install anthropic を実行してください。")

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("python-docx が未インストールです。pip install python-docx を実行してください。")

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

# ── 設定 ─────────────────────────────────────────────────────────────────────
CLAUDE_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
CLAUDE_MODEL   = "claude-sonnet-4-6"
CV_DIR         = Path(r"C:\Users\shondo\Desktop\agent_project\cv")
DB_BASE        = Path(r"C:\Users\shondo\Desktop\agent_project\db")
F              = 12  # 統一フォントサイズ (pt)

HEADERS        = {"User-Agent": "Mozilla/5.0 (compatible; CVMaker/1.0)"}
CHARS_PER_FILE = 4000  # DB ファイル 1 件あたりの最大読み込み文字数

# db/ 配下のフォルダー一覧と説明（Claude がフォルダーを選ぶための手がかり）
DB_FOLDERS = {
    "asean_oceania": "ASEAN諸国・オセアニア（タイ、ベトナム、インドネシア、シンガポール、豪州など）",
    "china":         "中国（習近平政権、中国共産党、人民元、香港など）",
    "europe":        "欧州（EU、英独仏伊、ウクライナ、ロシア、NATO欧州側など）",
    "geopolitics":   "地政学・地経学全般（半導体、サプライチェーン、鉱物資源など）",
    "india_mena":    "インド・中東・アフリカ（サウジ、イスラエル、イラン、UAE、ケニアなど）",
    "japan":         "日本（日本政府、日本企業、日本の外交・経済政策など）",
    "minerals":      "鉱物資源（レアアース、ニッケル、リチウム、黒鉛など）",
    "north america": "北米（米国、トランプ政権、カナダ、メキシコ、FRBなど）",
    "south america": "南米（ブラジル、アルゼンチン、ベネズエラなど）",
    "taiwan":        "台湾（台湾政府、台湾海峡、TSMC、台湾の安全保障など）",
    "trade_tariff_ai": "通商・関税・AI（WTO、FTA、人権、環境、AI規制など）",
    "others":        "上記以外",
}

# ── Claude プロンプト ──────────────────────────────────────────────────────────
INFO_PROMPT = """\
あなたは政治家・外交官・有識者の情報を整理する専門家です。
以下の人物について、日本の外交官が面談準備に使う略歴書用の情報を
JSON 形式で提供してください。

【対象人物】
{name}

【出力ルール】
- JSON のみを出力（前置き・説明・コードブロック記号は一切不要）
- 各フィールドの文字数は必ず守ること（Word 文書の 1 行に収める制約）
- 情報が不確かな場合は "不明" と記載

{{
  "full_name_ja": "日本語・カタカナ表記の氏名",
  "birth_date": "生年月日（例: 1957年 4月 22日）",
  "nationality": "国籍（例: ポーランド共和国）",
  "current_position": "現職（30文字以内）",
  "education": [
    {{"period": "在籍期間（例: 1975年 ～ 1980年）",
      "description": "学校名・学部（25文字以内）"}}
  ],
  "career": [
    {{"year": "年（例: 2007年）",
      "description": "役職・出来事（25文字以内、現職・功績に関係するものに絞る、最大 8 件）"}}
  ],
  "interests": ["関心事項・主要政策（各28文字以内、5〜7件）"],
  "todo": ["面談で歓迎される話題・姿勢（各28文字以内、5〜7件）"],
  "not_todo": ["面談で避けるべき言動（各28文字以内、5〜7件）"],
  "wikipedia_title_en": "英語版 Wikipedia の記事タイトル（例: Donald Tusk）"
}}
"""

FOLDER_DETECT_PROMPT = """\
以下の人物が「最も関係する地域・テーマ」に対応するフォルダーを1つだけ選んでください。

【人物情報】
氏名: {name}
国籍: {nationality}
現職: {current_position}

【選択肢】
{folder_list}

回答はフォルダー名のみ1行で返してください（説明不要）。
"""

OFFICIAL_PHOTO_PROMPT = """\
以下の人物の「現職」に関連する公式ウェブサイトに、
プロフィール写真またはポートレート写真が掲載されている可能性があります。

【人物情報】
氏名: {name}
現職: {current_position}
国籍: {nationality}

直接ダウンロード可能な画像ファイル URL を 1 件だけ提示してください。
優先順位:
  1. 大使・外交官の場合 → 派遣先大使館の公式サイト（例: jp.usembassy.gov 等）
  2. 政府高官の場合    → 所属省庁・政府機関の公式サイト
  3. 国際機関職員     → 当該機関の公式サイト

条件:
  - 直接アクセスできる画像ファイル（.jpg / .png / .webp など）の URL のみ
  - 存在しない・不明な場合は「不明」とだけ回答

回答は URL のみ 1 行で返してください（説明不要）。
"""

ENRICH_PROMPT = """\
あなたは外交官向け面談準備資料の専門家です。
以下の人物の略歴書草案に、データベースファイルの最新情報を反映して
interests・todo・not_todo を更新してください。

【対象人物】
氏名: {name}（{nationality}、{current_position}）

【データベースファイルの内容】
{files_block}

【現在の草案（interests / todo / not_todo）】
{current_sections}

【指示】
- データベースの情報を踏まえ、より具体的・最新の内容に更新する
- 草案の内容も活かしつつ、DB 情報で補強・修正する
- 各項目は28文字以内
- interests: 5〜7件、todo: 5〜7件、not_todo: 5〜7件
- JSON のみを出力（前置き不要）。形式:
{{"interests": [...], "todo": [...], "not_todo": [...]}}
"""


# ── XML / フォントヘルパー ─────────────────────────────────────────────────────
def _set_east_asian_font(run, font_name: str):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "FFFFFF")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def set_cell_width(cell, width_cm: float):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")


def set_cell_font(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(F)
    run.font.bold = bold
    run.font.name = "MS明朝"
    _set_east_asian_font(run, "MS明朝")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


# ── 文書生成ヘルパー ──────────────────────────────────────────────────────────
def para(doc, text: str = "", size: int = F, bold: bool = False,
         align=WD_ALIGN_PARAGRAPH.LEFT, color=None,
         space_after: int = 3, underline: bool = False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.size      = Pt(size)
        run.font.bold      = bold
        run.font.underline = underline
        run.font.name      = "MS明朝"
        _set_east_asian_font(run, "MS明朝")
        if color:
            run.font.color.rgb = color
    return p


def add_info_line(cell, label: str, value: str, first: bool = False):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(label)
    r1.font.size = Pt(F)
    r1.font.name = "MS明朝"
    _set_east_asian_font(r1, "MS明朝")
    r2 = p.add_run(f"　{value}")
    r2.font.size = Pt(F)
    r2.font.bold = True
    r2.font.name = "MS明朝"
    _set_east_asian_font(r2, "MS明朝")


def add_bullets(doc, items: list):
    for item in items:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after  = Pt(3)
        run = bp.add_run(item)
        run.font.size = Pt(F)
        run.font.name = "MS明朝"
        _set_east_asian_font(run, "MS明朝")


# ── Claude で人物情報を取得 ────────────────────────────────────────────────────
def fetch_person_info(name: str, client) -> dict:
    msg = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=2000,
        messages=[{"role": "user", "content": INFO_PROMPT.format(name=name)}],
    )
    raw = msg.content[0].text.strip()
    raw = raw.lstrip("﻿")  # BOM 除去
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s*```\s*$", "", raw)
    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"  JSON パースエラー: {e}")
        print(f"  出力の先頭 300 文字:\n{raw[:300]}")
        raise


# ── DB フォルダーの判定 ────────────────────────────────────────────────────────
def detect_db_folder(info: dict, client) -> Path | None:
    """人物情報をもとに最も関連する db/ サブフォルダーを返す。"""
    available = {k: v for k, v in DB_FOLDERS.items()
                 if (DB_BASE / k).is_dir()}
    if not available:
        return None

    folder_list = "\n".join(f"  {k}: {v}" for k, v in available.items())
    prompt = FOLDER_DETECT_PROMPT.format(
        name=info.get("full_name_ja", ""),
        nationality=info.get("nationality", ""),
        current_position=info.get("current_position", ""),
        folder_list=folder_list,
    )
    msg = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=30,
        messages=[{"role": "user", "content": prompt}],
    )
    chosen = msg.content[0].text.strip().lower().strip('"').strip("'")
    # 前後の空白や改行を除去し、選択肢に含まれるか確認
    for key in available:
        if key in chosen:
            folder = DB_BASE / key
            print(f"  [DB] 関連フォルダー: {folder}")
            return folder
    return None


# ── DB ファイルのテキスト抽出 ─────────────────────────────────────────────────
def _extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            if pdfplumber:
                with pdfplumber.open(str(path)) as pdf:
                    return "\n".join(p.extract_text() or "" for p in pdf.pages)
            if fitz:
                doc = fitz.open(str(path))
                return "\n".join(page.get_text() for page in doc)
            return ""
        if ext == ".docx":
            doc = Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        if ext in (".txt", ".md"):
            return path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        pass
    return ""


def build_files_block(folder: Path) -> str:
    """フォルダー内の対応ファイルをすべて読み込み、テキストブロックを返す。"""
    supported = {".pdf", ".docx", ".txt", ".md"}
    files = sorted(folder.rglob("*"))
    blocks = []
    for f in files:
        if f.suffix.lower() not in supported:
            continue
        text = _extract_text(f)
        if not text.strip():
            continue
        snippet = text[:CHARS_PER_FILE]
        if len(text) > CHARS_PER_FILE:
            snippet += "\n[以下省略]"
        blocks.append(f"=== {f.name} ===\n{snippet}")
        if len(blocks) >= 15:  # 最大 15 ファイル
            break
    return "\n\n".join(blocks)


# ── DB 情報で interests / todo / not_todo を補強 ──────────────────────────────
def enrich_with_db(info: dict, files_block: str, client) -> dict:
    """DB ファイルの内容を使って interests / todo / not_todo を更新する。"""
    current_sections = json.dumps({
        "interests": info.get("interests", []),
        "todo":      info.get("todo", []),
        "not_todo":  info.get("not_todo", []),
    }, ensure_ascii=False, indent=2)

    prompt = ENRICH_PROMPT.format(
        name=info.get("full_name_ja", ""),
        nationality=info.get("nationality", ""),
        current_position=info.get("current_position", ""),
        files_block=files_block,
        current_sections=current_sections,
    )
    msg = client.messages.create(
        model=CLAUDE_MODEL,
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}],
    )
    raw = msg.content[0].text.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s*```\s*$", "", raw)
    try:
        enriched = json.loads(raw)
        for key in ("interests", "todo", "not_todo"):
            if key in enriched and isinstance(enriched[key], list):
                info[key] = enriched[key]
        return info
    except json.JSONDecodeError as e:
        print(f"  [DB補強] JSON パースエラー: {e}  草案のまま使用します。")
        return info


# ── Wikipedia から写真を取得 ──────────────────────────────────────────────────
def fetch_wikipedia_photo(wikipedia_title: str, save_path: Path) -> bool:
    if not wikipedia_title or wikipedia_title == "不明":
        return False
    try:
        encoded = urllib.parse.quote(wikipedia_title.replace(" ", "_"))
        api_url = (
            "https://en.wikipedia.org/w/api.php"
            f"?action=query&titles={encoded}&prop=pageimages"
            "&format=json&pithumbsize=400"
        )
        req = urllib.request.Request(api_url, headers=HEADERS)
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = json.loads(resp.read())

        pages = data.get("query", {}).get("pages", {})
        thumb_url = None
        for page in pages.values():
            thumb_url = (page.get("thumbnail") or {}).get("source")
            break

        if not thumb_url:
            print("  [写真] Wikipedia に写真が見つかりませんでした。写真なしで作成します。")
            return False

        print("  [写真] 取得中 ...")
        req2 = urllib.request.Request(thumb_url, headers=HEADERS)
        with urllib.request.urlopen(req2, timeout=10) as resp:
            img_data = resp.read()

        if Image:
            img = Image.open(BytesIO(img_data)).convert("RGB")
            img.save(str(save_path), "JPEG", quality=90)
        else:
            save_path.write_bytes(img_data)

        print(f"  [写真] 保存しました: {save_path.name}")
        return True
    except Exception as e:
        print(f"  [写真] 取得エラー: {e}  写真なしで作成します。")
        return False


# ── 公式サイトから写真を取得 ──────────────────────────────────────────────────
def fetch_official_photo(info: dict, client, save_path: Path) -> bool:
    """役職・国籍をもとに公式サイトの写真URLをClaudeに推定させ、取得する。"""
    prompt = OFFICIAL_PHOTO_PROMPT.format(
        name=info.get("full_name_ja", ""),
        current_position=info.get("current_position", ""),
        nationality=info.get("nationality", ""),
    )
    try:
        msg = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=200,
            messages=[{"role": "user", "content": prompt}],
        )
        url = msg.content[0].text.strip()
    except Exception as e:
        print(f"  [公式写真] URL取得エラー: {e}")
        return False

    if not url or url == "不明" or not url.startswith("http"):
        print("  [公式写真] 公式サイトの写真URLが見つかりませんでした。")
        return False

    print(f"  [公式写真] 取得中: {url}")
    try:
        req = urllib.request.Request(url, headers=HEADERS)
        with urllib.request.urlopen(req, timeout=10) as resp:
            img_data = resp.read()

        if Image:
            img = Image.open(BytesIO(img_data)).convert("RGB")
            img.save(str(save_path), "JPEG", quality=90)
        else:
            save_path.write_bytes(img_data)

        print(f"  [公式写真] 保存しました: {save_path.name}")
        return True
    except Exception as e:
        print(f"  [公式写真] 取得エラー: {e}")
        return False


# ── Word 文書を生成 ───────────────────────────────────────────────────────────
def generate_doc(info: dict, photo_path, output_path: Path):
    doc = Document()

    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.0)

    # ── タイトル
    para(doc, "履　　　歴　　　書", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    # ── 基本情報 + 写真
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style   = "Table Grid"
    tbl.autofit = False
    left  = tbl.rows[0].cells[0]
    right = tbl.rows[0].cells[1]
    set_cell_width(left,  12.8)
    set_cell_width(right,  4.2)

    add_info_line(left, "氏　　名", info.get("full_name_ja", ""), first=True)
    add_info_line(left, "生年月日", info.get("birth_date", ""))
    add_info_line(left, "国　　籍", info.get("nationality", ""))
    add_info_line(left, "現　　職", info.get("current_position", ""))
    remove_table_borders(tbl)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(0)
    if photo_path and Path(photo_path).exists():
        rp.add_run().add_picture(str(photo_path), width=Cm(4.0))

    doc.add_paragraph()

    # ── 学歴
    para(doc, "学　　　　　歴", size=F, bold=True, space_after=2, underline=True)
    edu_list = info.get("education", [])
    if edu_list:
        edu = doc.add_table(rows=len(edu_list), cols=2)
        edu.style   = "Table Grid"
        edu.autofit = False
        for i, e in enumerate(edu_list):
            set_cell_font(edu.rows[i].cells[0], e.get("period", ""))
            set_cell_font(edu.rows[i].cells[1], e.get("description", ""))
            set_cell_width(edu.rows[i].cells[0], 4.7)
            set_cell_width(edu.rows[i].cells[1], 11.8)
        remove_table_borders(edu)

    doc.add_paragraph()

    # ── 主な職歴
    para(doc, "主　な　職　歴", size=F, bold=True, space_after=2, underline=True)
    career_list = info.get("career", [])
    if career_list:
        ctbl = doc.add_table(rows=len(career_list), cols=2)
        ctbl.style   = "Table Grid"
        ctbl.autofit = False
        for i, c in enumerate(career_list):
            set_cell_font(ctbl.rows[i].cells[0], c.get("year", ""))
            set_cell_font(ctbl.rows[i].cells[1], c.get("description", ""))
            set_cell_width(ctbl.rows[i].cells[0], 4.7)
            set_cell_width(ctbl.rows[i].cells[1], 11.8)
        remove_table_borders(ctbl)

    doc.add_paragraph()

    # ── 関心事項・主要政策
    para(doc, "関　心　事　項・主　要　政　策", size=F, bold=True,
         space_after=4, underline=True)
    add_bullets(doc, info.get("interests", []))
    doc.add_paragraph()

    # ── To Do
    para(doc, "To Do　（歓迎される話題・姿勢）", size=F, bold=True,
         space_after=4, underline=True)
    add_bullets(doc, info.get("todo", []))
    doc.add_paragraph()

    # ── Not to Do
    para(doc, "Not to Do　（避けるべき言動）", size=F, bold=True,
         space_after=4, underline=True)
    add_bullets(doc, info.get("not_todo", []))
    doc.add_paragraph()

    # ── 作成日
    date_str = datetime.now().strftime("%Y年　%m月　%d日　作成")
    para(doc, date_str, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT,
         color=RGBColor(0x60, 0x60, 0x60), space_after=0)

    doc.save(str(output_path))


# ── メイン ────────────────────────────────────────────────────────────────────
def main():
    CV_DIR.mkdir(parents=True, exist_ok=True)

    print("=" * 50)
    print("  略歴書ジェネレーター")
    print("=" * 50)
    try:
        name = input(
            "\n誰の略歴書を作成しますか？\n"
            "（例: マクロン仏大統領、習近平、石破茂）\n> "
        ).strip()
    except (EOFError, KeyboardInterrupt):
        print("\nキャンセルされました。")
        return

    if not name:
        sys.exit("名前が入力されませんでした。")

    print(f"\n  [Claude] 情報を問い合わせ中: {name} ...")
    sys.stdout.flush()

    client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)

    # 1. 人物情報を Claude で取得
    try:
        info = fetch_person_info(name, client)
    except Exception as e:
        sys.exit(f"情報取得に失敗しました: {e}")

    full_name_ja = info.get("full_name_ja", name)
    print(f"\n  氏名: {full_name_ja}")
    print(f"  国籍: {info.get('nationality', '?')}")
    print(f"  現職: {info.get('current_position', '?')}")

    # 2. DB フォルダーから情報を取得して補強
    print("\n  [DB] 関連フォルダーを判定中...")
    db_folder = detect_db_folder(info, client)
    if db_folder:
        print(f"  [DB] ファイルを読み込み中...")
        files_block = build_files_block(db_folder)
        if files_block:
            file_count = files_block.count("=== ")
            print(f"  [DB] {file_count} 件のファイルから情報を取得。Claude で分析補強中...")
            info = enrich_with_db(info, files_block, client)
            print("  [DB] 補強完了。")
        else:
            print("  [DB] 読み込めるファイルがありませんでした。")
    else:
        print("  [DB] 該当フォルダーなし。スキップします。")

    # 3. 写真を取得（公式サイト優先 → Wikipedia フォールバック）
    safe_name  = re.sub(r'[\\/:*?"<>|\s]', "_", full_name_ja)
    photo_dir  = CV_DIR / "photo"
    photo_dir.mkdir(parents=True, exist_ok=True)
    photo_path = photo_dir / f"{safe_name}.jpg"

    print("\n  [写真] 公式サイトから取得を試みています...")
    has_photo = fetch_official_photo(info, client, photo_path)
    if not has_photo:
        print("  [写真] Wikipedia から取得を試みています...")
        has_photo = fetch_wikipedia_photo(info.get("wikipedia_title_en", ""), photo_path)
    if not has_photo:
        photo_path = None

    # 4. Word 文書を生成
    date_tag    = datetime.now().strftime("%Y%m%d")
    output_path = CV_DIR / f"{safe_name}_略歴書_{date_tag}.docx"

    print("\n  [Word] 文書を生成中...")
    try:
        generate_doc(info, photo_path, output_path)
        print(f"\n保存完了: {output_path}")
    except Exception as e:
        sys.exit(f"文書生成に失敗しました: {e}")


if __name__ == "__main__":
    main()
